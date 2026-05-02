"""
Dịch vụ tạo bìa chứng từ - docxtpl (template-based)

Template: templates/bia_mau_goc.docx
Placeholders: dept_name, date_text, tap_so, total_sheets, custodian,
              r{0-7}_{lu,ln,ld,lc,ru,rn,rd,rc}
"""
import io
import os
import copy
from typing import List
from datetime import date

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.services.bundle_service import BundleResult, bundle_label

TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "templates", "bia_mau_goc.docx",
)


def _format_date_for_header(dates: List[date]) -> str:
    if not dates:
        return "Ngày ... tháng ... năm ..."
    dates_sorted = sorted(set(dates))
    years = {d.year for d in dates_sorted}
    months = {d.month for d in dates_sorted}
    if len(years) == 1 and len(months) == 1:
        days = ", ".join(str(d.day) for d in dates_sorted)
        return f"Ngày {days} tháng {dates_sorted[0].month:02d} năm {dates_sorted[0].year}"
    elif len(years) == 1:
        parts = ", ".join(f"{d.day}/{d.month}" for d in dates_sorted)
        return f"Ngày {parts} năm {dates_sorted[0].year}"
    else:
        parts = ", ".join(f"{d.day}/{d.month}/{d.year}" for d in dates_sorted)
        return f"Ngày {parts}"


def _build_context(department_name: str, bundle: BundleResult) -> dict:
    units_sorted = sorted(bundle.units, key=lambda u: (u.transaction_date, u.is_large, u.user_code))
    left_col = units_sorted[:8]
    right_col = units_sorted[8:16]

    dates = sorted({u.transaction_date for u in bundle.units})
    ctx = {
        "dept_name": department_name.upper(),
        "date_text": _format_date_for_header(dates),
        "tap_so": bundle_label(bundle.label_seq, bundle.label_total),
        "total_sheets": str(bundle.total_sheets),
        "custodian": bundle.custodian_name,
    }

    for n in range(8):
        left = left_col[n] if n < len(left_col) else None
        right = right_col[n] if n < len(right_col) else None

        if left:
            ctx[f"r{n}_lu"] = left.user_code
            ctx[f"r{n}_ln"] = left.full_name or ""
            ctx[f"r{n}_ld"] = left.transaction_date.strftime("%d/%m/%Y")
            ctx[f"r{n}_lc"] = str(left.sheet_count)
        else:
            ctx[f"r{n}_lu"] = ctx[f"r{n}_ln"] = ctx[f"r{n}_ld"] = ctx[f"r{n}_lc"] = ""

        if right:
            ctx[f"r{n}_ru"] = right.user_code
            ctx[f"r{n}_rn"] = right.full_name or ""
            ctx[f"r{n}_rd"] = right.transaction_date.strftime("%d/%m/%Y")
            ctx[f"r{n}_rc"] = str(right.sheet_count)
        else:
            ctx[f"r{n}_ru"] = ctx[f"r{n}_rn"] = ctx[f"r{n}_rd"] = ctx[f"r{n}_rc"] = ""

    return ctx


def _render_to_doc(ctx: dict) -> Document:
    tpl = DocxTemplate(TEMPLATE_PATH)
    tpl.render(ctx)
    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    doc = Document(buf)
    # Center-align count cells (cells that contain only digits) in data rows
    if doc.tables:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        tbl = doc.tables[0]
        for row in tbl.rows[2:10]:
            for cell in row.cells:
                if cell.text.strip().isdigit():
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def generate_covers_docx(
    department_name: str,
    bundles: List[BundleResult],
) -> bytes:
    if not bundles:
        buf = io.BytesIO()
        Document().save(buf)
        return buf.getvalue()

    rendered = [_render_to_doc(_build_context(department_name, b)) for b in bundles]

    if len(rendered) == 1:
        buf = io.BytesIO()
        rendered[0].save(buf)
        return buf.getvalue()

    # Combine: keep base doc (with its page/style settings), append tables from others.
    #
    # The rendered body structure is: <tbl> <p> <sectPr>
    # A full-page table pushes the trailing <p> to a new page, so using a plain
    # page-break run inside it creates a blank page.  The correct OOXML fix is to
    # convert that trailing <p> into a section-break paragraph (w:sectPr inside
    # w:pPr with w:type="nextPage").  This paragraph fulfils the mandatory
    # "cursor-after-table" requirement AND transitions to the next page — no blanks.

    base = rendered[0]
    base_body = base._body._body

    # Use the document-level sectPr as the template for section-break paragraphs
    # (copies page size, margins etc. already set by the cover template).
    doc_sect_pr = base_body.find(qn("w:sectPr"))

    def _make_next_page_sect_pr() -> OxmlElement:
        sect = copy.deepcopy(doc_sect_pr)
        existing = sect.find(qn("w:type"))
        if existing is not None:
            sect.remove(existing)
        t = OxmlElement("w:type")
        t.set(qn("w:val"), "nextPage")
        sect.insert(0, t)
        return sect

    for doc in rendered[1:]:
        # Convert the trailing <p> into a section-break paragraph
        last_p = None
        for child in reversed(list(base_body)):
            if child.tag == qn("w:p"):
                last_p = child
                break

        if last_p is not None:
            pPr = last_p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                last_p.insert(0, pPr)
            pPr.append(_make_next_page_sect_pr())
        else:
            p_el = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            pPr.append(_make_next_page_sect_pr())
            p_el.append(pPr)
            current_sect = base_body.find(qn("w:sectPr"))
            if current_sect is not None:
                current_sect.addprevious(p_el)
            else:
                base_body.append(p_el)

        # Append next cover's table, then its trailing <p>
        current_sect = base_body.find(qn("w:sectPr"))
        for tbl in doc.tables:
            tbl_copy = copy.deepcopy(tbl._tbl)
            current_sect = base_body.find(qn("w:sectPr"))
            current_sect.addprevious(tbl_copy)

        # Carry the trailing <p> (clean, no section break — may be promoted in next loop)
        for child in doc._body._body:
            if child.tag == qn("w:p"):
                current_sect = base_body.find(qn("w:sectPr"))
                current_sect.addprevious(copy.deepcopy(child))
                break

    buf = io.BytesIO()
    base.save(buf)
    return buf.getvalue()
