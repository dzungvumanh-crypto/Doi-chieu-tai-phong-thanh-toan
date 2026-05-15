"""Xử lý file IPCAS / Payment → sinh báo cáo hậu kiểm"""
import io
import os
import sqlite3
import zipfile
from pathlib import Path
import pandas as pd
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

_TEMPLATES_DIR = Path(__file__).parents[2] / 'templates'


# ── Helpers ───────────────────────────────────────────────────────────────────

def _to_int(val) -> int:
    try:
        v = str(val).strip().replace(',', '')
        if v.lower() in ('', 'nan', 'none', '-'):
            return 0
        return int(float(v))
    except Exception:
        return 0


def _get_month_from_date(date_str: str) -> int:
    try:
        parts = str(date_str).strip().split('/')
        if len(parts) >= 2:
            return int(parts[1])
    except Exception:
        pass
    return 0


def _get_year_from_date(date_str: str) -> int:
    try:
        parts = str(date_str).strip().split('/')
        if len(parts) >= 3:
            return int(parts[2][:4])
    except Exception:
        pass
    return 0


# ── Parse IPCAS (GDV hoặc HKV) ───────────────────────────────────────────────

def parse_ipcas(file_bytes: bytes) -> tuple:
    df = pd.read_excel(io.BytesIO(file_bytes), engine='xlrd', header=0)
    violations = []
    month, year = 0, 0
    aggregated: dict = {}

    for df_idx, row in df.iterrows():
        user_code = str(row.get('col_tellerid', '')).strip()
        if not user_code or user_code.lower() in ('nan', 'col_tellerid', ''):
            continue
        gd_sai = _to_int(row.get('name_5', 0))
        if gd_sai != 0:
            violations.append({'user_code': user_code, 'name_5': gd_sai, 'row_num': int(df_idx) + 2})

        date_val = str(row.get('name_1', '')).strip()
        if not month:
            month = _get_month_from_date(date_val)
            year = _get_year_from_date(date_val)

        if user_code not in aggregated:
            aggregated[user_code] = {'user_code': user_code, 'tong_gd': 0, 'gd_hk_dung': 0,
                                     'gd_sai': 0, 'chua_hk': 0, 'bt_huy': 0}
        agg = aggregated[user_code]
        agg['tong_gd']    += _to_int(row.get('name_3', 0))
        agg['gd_hk_dung'] += _to_int(row.get('name_4', 0))
        agg['gd_sai']     += gd_sai
        agg['chua_hk']    += _to_int(row.get('name_7', 0))
        agg['bt_huy']     += _to_int(row.get('name_10', 0))

    rows = list(aggregated.values())
    return rows, violations, month, year


# ── Parse Payment (Teller hoặc Backchecker) ───────────────────────────────────

def _parse_payment(file_bytes: bytes, header_row: int) -> tuple:
    df = pd.read_excel(io.BytesIO(file_bytes), header=header_row)
    df = df[df['STT'].apply(lambda x: _to_int(x) > 0)].copy()

    col_map = {
        'tong_gd': 'Tổng GD',
        'gd_hk_dung': 'GD Hậu kiểm đúng',
        'gd_sai': 'GD Hậu kiểm sai',
        'chua_hk': 'GD chưa HK',
        'bt_huy': 'Tổng số GD Hủy',
    }

    rows, violations = [], []
    for hkv_username, grp in df.groupby('HKV'):
        username = str(hkv_username).strip()
        if not username or username.lower() == 'nan':
            continue
        totals = {}
        for key, col in col_map.items():
            try:
                totals[key] = int(grp[col].apply(_to_int).sum())
            except Exception:
                totals[key] = 0

        if totals.get('gd_sai', 0) != 0:
            violations.append({'user_code': username, 'name_5': totals['gd_sai']})

        rows.append({'user_code': username, **totals})

    return rows, violations


def parse_payment_teller(file_bytes: bytes) -> tuple:
    return _parse_payment(file_bytes, header_row=6)


def parse_payment_backchecker(file_bytes: bytes) -> tuple:
    return _parse_payment(file_bytes, header_row=10)


# ── Enrich rows → gắn tên + phòng ────────────────────────────────────────────

def enrich_and_group(rows: list, db: sqlite3.Connection, is_payment: bool = False) -> dict:
    codes = [r['user_code'] for r in rows]
    if not codes:
        return {}

    lower_codes = [c.strip().lower() for c in codes]
    ph = ",".join("?" * len(lower_codes))

    if is_payment:
        users_rows = db.execute(
            f"""SELECT ks.*, d.name AS dept_name
                FROM ksnb_staff ks
                LEFT JOIN departments d ON ks.department_id = d.id
                WHERE lower(trim(ks.payment_username)) IN ({ph})""",
            lower_codes,
        ).fetchall()
        lookup = {(u["payment_username"] or '').strip().lower(): dict(u) for u in users_rows}

        # Suffix fallback: thaonguyenthiphuong → thaonguyenthiphuong1
        for code in lower_codes:
            if code not in lookup:
                for suffix in ('1', '2', '3'):
                    if code + suffix in lookup:
                        lookup[code] = lookup[code + suffix]
                        break
                if code not in lookup:
                    row = db.execute(
                        """SELECT ks.*, d.name AS dept_name
                           FROM ksnb_staff ks
                           LEFT JOIN departments d ON ks.department_id = d.id
                           WHERE lower(trim(ks.payment_username)) LIKE ?""",
                        (code + '%',),
                    ).fetchone()
                    if row:
                        lookup[code] = dict(row)
    else:
        users_rows = db.execute(
            f"""SELECT ks.*, d.name AS dept_name
                FROM ksnb_staff ks
                LEFT JOIN departments d ON ks.department_id = d.id
                WHERE lower(trim(ks.ipcas_code)) IN ({ph})""",
            lower_codes,
        ).fetchall()
        lookup = {(u["ipcas_code"] or '').strip().lower(): dict(u) for u in users_rows}

    by_dept: dict = {}
    for row in rows:
        lookup_key = row['user_code'].strip().lower()
        user = lookup.get(lookup_key)
        if user:
            vn_name = user.get("full_name") or user.get("payment_username") or row['user_code']
            user_code_ipcas = user.get("ipcas_code")
            dept_name = user.get("dept_name") or 'Không xác định'
        else:
            vn_name = row['user_code']
            user_code_ipcas = row['user_code']
            dept_name = 'Không xác định'

        enriched = {
            **row,
            'vn_name': vn_name,
            'user_code_display': user_code_ipcas,
            'dept_name': dept_name,
            'hkv_thuc_hien': '',
            'huy_kq': 0,
            'huy_cq': 0,
            'nguyen_nhan': '',
        }
        by_dept.setdefault(dept_name, []).append(enriched)

    return by_dept


# ── Merge IPCAS + Payment HKV theo user ──────────────────────────────────────

def merge_hkv(ipcas_rows: list, payment_rows: list, db: sqlite3.Connection) -> list:
    ipcas_by_code = {r['user_code']: r for r in ipcas_rows}
    payment_by_name = {r['user_code'].lower(): r for r in payment_rows}

    ipcas_codes = list(ipcas_by_code.keys())
    payment_names = list(payment_by_name.keys())

    ipcas_users: dict = {}
    if ipcas_codes:
        lower_ipcas = [c.strip().lower() for c in ipcas_codes]
        ph = ",".join("?" * len(lower_ipcas))
        for u in db.execute(
            f"""SELECT ks.*, d.name AS dept_name
                FROM ksnb_staff ks
                LEFT JOIN departments d ON ks.department_id = d.id
                WHERE lower(trim(ks.ipcas_code)) IN ({ph})""",
            lower_ipcas,
        ).fetchall():
            ipcas_users[(u["ipcas_code"] or '').strip().lower()] = dict(u)

    payment_users: dict = {}
    if payment_names:
        ph = ",".join("?" * len(payment_names))
        for u in db.execute(
            f"""SELECT ks.*, d.name AS dept_name
                FROM ksnb_staff ks
                LEFT JOIN departments d ON ks.department_id = d.id
                WHERE lower(trim(ks.payment_username)) IN ({ph})""",
            payment_names,
        ).fetchall():
            payment_users[(u["payment_username"] or '').strip().lower()] = dict(u)

    merged = []
    seen_ipcas = set()

    for code, ipcas_row in ipcas_by_code.items():
        u = ipcas_users.get(code.lower())
        vn_name = (u.get("full_name") or u.get("payment_username") or code) if u else code
        payment_username = u.get("payment_username") if u else None
        payment_row = payment_by_name.get(payment_username.lower()) if payment_username else None

        merged.append({
            'user_code': code,
            'vn_name': vn_name,
            'ipcas_tong_gd': ipcas_row['tong_gd'],
            'ipcas_gd_hk_dung': ipcas_row['gd_hk_dung'],
            'ipcas_gd_sai': ipcas_row.get('gd_sai', 0),
            'ipcas_chua_hk': ipcas_row['chua_hk'],
            'ipcas_bt_huy': ipcas_row['bt_huy'],
            'payment_tong_gd': payment_row['tong_gd'] if payment_row else 0,
            'payment_gd_hk_dung': payment_row['gd_hk_dung'] if payment_row else 0,
            'payment_gd_sai': (payment_row.get('gd_sai', 0) if payment_row else 0),
            'payment_chua_hk': payment_row['chua_hk'] if payment_row else 0,
            'payment_bt_huy': payment_row['bt_huy'] if payment_row else 0,
        })
        seen_ipcas.add((u.get("ipcas_code") or code).lower() if u else code.lower())

    for username, payment_row in payment_by_name.items():
        u = payment_users.get(username)
        if u and (u.get("ipcas_code") or '').lower() in seen_ipcas:
            continue
        vn_name = (u.get("full_name") or u.get("payment_username") or username) if u else username
        merged.append({
            'user_code': u.get("ipcas_code") if u else username,
            'vn_name': vn_name,
            'ipcas_tong_gd': 0,
            'ipcas_gd_hk_dung': 0,
            'ipcas_gd_sai': 0,
            'ipcas_chua_hk': 0,
            'ipcas_bt_huy': 0,
            'payment_tong_gd': payment_row['tong_gd'],
            'payment_gd_hk_dung': payment_row['gd_hk_dung'],
            'payment_gd_sai': payment_row.get('gd_sai', 0),
            'payment_chua_hk': payment_row['chua_hk'],
            'payment_bt_huy': payment_row['bt_huy'],
        })

    return merged


# ── Generate Excel báo cáo phòng ─────────────────────────────────────────────

def _save_row_styles(ws, row_num: int) -> tuple:
    from copy import copy
    styles = []
    for col in range(1, 12):
        cell = ws.cell(row_num, col)
        styles.append({
            'font':          copy(cell.font),
            'alignment':     copy(cell.alignment),
            'border':        copy(cell.border),
            'fill':          copy(cell.fill),
            'number_format': cell.number_format,
        })
    height = ws.row_dimensions[row_num].height
    return styles, height


def _apply_row_styles(ws, row_num: int, styles: list, height):
    from copy import copy
    for col, s in enumerate(styles, 1):
        cell = ws.cell(row_num, col)
        cell.font          = copy(s['font'])
        cell.alignment     = copy(s['alignment'])
        cell.border        = copy(s['border'])
        cell.fill          = copy(s['fill'])
        cell.number_format = s['number_format']
    if height:
        ws.row_dimensions[row_num].height = height


def generate_dept_excel(dept_name: str, month: int, year: int,
                        ipcas_rows: list, payment_rows: list,
                        staff_name: str) -> bytes:
    template_path = _TEMPLATES_DIR / 'Báo cáo Tổng hợp hậu kiểm theo phòng.xlsx'
    wb = openpyxl.load_workbook(str(template_path))
    ws = wb.active
    ws.title = f'T{month:02d}.{year}'

    ws['A5'].value = f'BÁO CÁO TỔNG HỢP HẬU KIỂM PHÒNG {dept_name.upper()}'
    ws['A6'].value = f'Tháng {month:02d} năm {year}'

    sec_styles,  sec_h  = _save_row_styles(ws, 8)
    data_styles, data_h = _save_row_styles(ws, 9)
    tot_styles,  tot_h  = _save_row_styles(ws, 16)
    sign_styles, sign_h = _save_row_styles(ws, 26)
    name_styles, name_h = _save_row_styles(ws, 30)

    to_unmerge = [str(m) for m in ws.merged_cells.ranges if m.min_row >= 8]
    for rng in to_unmerge:
        ws.unmerge_cells(rng)
    if ws.max_row >= 8:
        ws.delete_rows(8, ws.max_row - 7)

    def _write_data(r: int, stt: int, d: dict):
        _apply_row_styles(ws, r, data_styles, data_h)
        ws.cell(r, 1).value  = stt
        ws.cell(r, 2).value  = d.get('user_code_display', d.get('user_code', ''))
        ws.cell(r, 3).value  = d.get('vn_name', '')
        ws.cell(r, 4).value  = d.get('hkv_thuc_hien', '') or None
        ws.cell(r, 5).value  = d.get('gd_hk_dung', 0) or None
        ws.cell(r, 6).value  = d.get('chua_hk', 0) or None
        ws.cell(r, 7).value  = d.get('bt_huy', 0) or None
        ws.cell(r, 8).value  = d.get('huy_kq', 0) or None
        ws.cell(r, 9).value  = d.get('huy_cq', 0) or None
        ws.cell(r, 10).value = f'=IFERROR({ws.cell(r,9).coordinate}/{ws.cell(r,5).coordinate}*100,0)'
        ws.cell(r, 11).value = d.get('nguyen_nhan', '') or None

    def _write_total(r: int, label: str, start: int, end: int):
        _apply_row_styles(ws, r, tot_styles, tot_h)
        ws.cell(r, 2).value = label
        for ci, col in [(5,'E'),(6,'F'),(7,'G'),(8,'H'),(9,'I')]:
            ws.cell(r, ci).value = f'=SUM({col}{start}:{col}{end})'
        ws.cell(r, 10).value = f'=IFERROR({ws.cell(r,9).coordinate}/{ws.cell(r,5).coordinate}*100,0)'

    r = 8
    ws.merge_cells(f'A{r}:K{r}')
    _apply_row_styles(ws, r, sec_styles, sec_h)
    ws.cell(r, 1).value = 'I. HỆ THỐNG IPCAS'
    r += 1

    ipcas_start = r
    for stt, d in enumerate(ipcas_rows, 1):
        _write_data(r, stt, d); r += 1
    ipcas_end = r - 1
    _write_total(r, 'TC (I)', ipcas_start, ipcas_end); r += 1

    ws.merge_cells(f'A{r}:K{r}')
    _apply_row_styles(ws, r, sec_styles, sec_h)
    ws.cell(r, 1).value = 'II. HỆ THỐNG THANH TOÁN TẬP TRUNG'
    r += 1

    payment_start = r
    for stt, d in enumerate(payment_rows, 1):
        _write_data(r, stt, d); r += 1
    payment_end = r - 1
    _write_total(r, 'TC (II)', payment_start, payment_end); r += 2

    _apply_row_styles(ws, r, sign_styles, sign_h)
    ws.cell(r, 11).value = f'Hà Nội, ngày ... tháng {month:02d} năm {year}'
    r += 1
    _apply_row_styles(ws, r, sign_styles, sign_h)
    ws.cell(r, 11).value = 'Người làm báo cáo'
    r += 4
    _apply_row_styles(ws, r, name_styles, name_h)
    ws.cell(r, 11).value = staff_name

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ── Generate Word báo cáo tổng hợp trung tâm ─────────────────────────────────

def _replace_date_runs(para, month: int, year: int):
    import re as _re
    runs = para.runs
    if not runs:
        return

    full_text = ''.join(r.text for r in runs)
    replacements_spec = [
        (_re.compile(r'([Tt]háng\s+)\d{2}/\d{4}'),
         lambda m: f'{m.group(1)}{month:02d}/{year}'),
        (_re.compile(r'(THÁNG\s+)\d{2}(\s+NĂM\s+)\d{4}'),
         lambda m: f'{m.group(1)}{month:02d}{m.group(2)}{year}'),
    ]

    all_matches = []
    for pat, fn in replacements_spec:
        for m in pat.finditer(full_text):
            all_matches.append((m.start(), m.end(), fn(m)))
    if not all_matches:
        return
    all_matches.sort(key=lambda x: x[0], reverse=True)

    pos = 0
    spans = []
    for r in runs:
        spans.append((pos, pos + len(r.text)))
        pos += len(r.text)

    for match_start, match_end, replacement in all_matches:
        overlap = [(i, s, e) for i, (s, e) in enumerate(spans)
                   if s < match_end and e > match_start]
        if not overlap:
            continue
        fi, fs, _  = overlap[0]
        li, _,  le = overlap[-1]
        prefix = full_text[fs:match_start]
        suffix = full_text[match_end:le]

        if fi == li:
            runs[fi].text = prefix + replacement + suffix
        else:
            runs[fi].text = prefix + replacement
            for i, _, _ in overlap[1:-1]:
                runs[i].text = ''
            runs[li].text = suffix

        pos = 0
        spans = []
        full_text = ''.join(r.text for r in runs)
        for r in runs:
            spans.append((pos, pos + len(r.text)))
            pos += len(r.text)


def _set_row_cells(row, values):
    import copy as _copy
    seen = set()
    vi = 0
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        val = str(values[vi]) if vi < len(values) else ''
        vi += 1
        para = cell.paragraphs[0] if cell.paragraphs else None
        if para is None:
            cell.text = val
            continue
        if para.runs:
            para.runs[0].text = val
            for run in para.runs[1:]:
                run.text = ''
        else:
            para.add_run(val)


def generate_center_word(month: int, year: int,
                         hkv_ipcas: list, hkv_payment: list,
                         dept_data: list) -> bytes:
    import copy
    template_path = _TEMPLATES_DIR / 'Báo cáo hậu kiểm tháng.docx'
    doc = Document(str(template_path))

    for para in doc.paragraphs:
        _replace_date_runs(para, month, year)

    def _set_date_cell(cell, text: str):
        para = cell.paragraphs[-1] if cell.paragraphs else None
        if para is None:
            return
        for p in cell.paragraphs:
            if 'ngày' in p.text or 'Hà Nội' in p.text:
                para = p
                break
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ''
        else:
            para.add_run(text)

    date_str = f'Hà Nội, ngày .... tháng {month:02d} năm {year}'
    _set_date_cell(doc.tables[0].rows[1].cells[1], date_str)
    _set_date_cell(doc.tables[3].rows[0].cells[1], date_str)

    t1 = doc.tables[1]
    data_row_xml    = copy.deepcopy(t1.rows[3]._tr)
    total_i_xml     = copy.deepcopy(t1.rows[11]._tr)
    section_ii_xml  = copy.deepcopy(t1.rows[12]._tr)
    total_ii_xml    = copy.deepcopy(t1.rows[20]._tr)
    grand_total_xml = copy.deepcopy(t1.rows[21]._tr)

    while len(t1.rows) > 3:
        t1._tbl.remove(t1.rows[-1]._tr)

    def _append_data_row(table, xml_tpl, vals):
        table._tbl.append(copy.deepcopy(xml_tpl))
        _set_row_cells(table.rows[-1], vals)

    ti_gd = ti_dung = ti_sai = ti_chua = ti_huy = 0
    for stt, row in enumerate(hkv_ipcas, 1):
        tong = row.get('ipcas_tong_gd', 0); dung = row.get('ipcas_gd_hk_dung', 0)
        sai  = row.get('ipcas_gd_sai', 0);  chua = row.get('ipcas_chua_hk', 0)
        huy  = row.get('ipcas_bt_huy', 0)
        ty   = f"{round(dung/tong*100, 2):.2f}" if tong else "0.00"
        ti_gd += tong; ti_dung += dung; ti_sai += sai; ti_chua += chua; ti_huy += huy
        _append_data_row(t1, data_row_xml, [str(stt), '1000', row.get('user_code', ''),
                                            str(tong), str(dung), str(sai), str(chua), str(huy), ty])
    ti_ty = f"{round(ti_dung/ti_gd*100, 2):.2f}" if ti_gd else "0.00"
    _append_data_row(t1, total_i_xml, ['', 'Tổng (I):', str(ti_gd), str(ti_dung),
                                       str(ti_sai), str(ti_chua), str(ti_huy), ti_ty])
    t1._tbl.append(copy.deepcopy(section_ii_xml))

    tp_gd = tp_dung = tp_sai = tp_chua = tp_huy = 0
    for stt, row in enumerate(hkv_payment, 1):
        tong = row.get('payment_tong_gd', 0); dung = row.get('payment_gd_hk_dung', 0)
        sai  = row.get('payment_gd_sai', 0);  chua = row.get('payment_chua_hk', 0)
        huy  = row.get('payment_bt_huy', 0)
        ty   = f"{round(dung/tong*100, 2):.2f}" if tong else "0.00"
        tp_gd += tong; tp_dung += dung; tp_sai += sai; tp_chua += chua; tp_huy += huy
        _append_data_row(t1, data_row_xml, [str(stt), '1000', row.get('user_code', ''),
                                            str(tong), str(dung), str(sai), str(chua), str(huy), ty])
    tp_ty = f"{round(tp_dung/tp_gd*100, 2):.2f}" if tp_gd else "0.00"
    _append_data_row(t1, total_ii_xml, ['', 'Tổng (II):', str(tp_gd), str(tp_dung),
                                        str(tp_sai), str(tp_chua), str(tp_huy), tp_ty])
    gr_gd = ti_gd + tp_gd; gr_dung = ti_dung + tp_dung
    gr_sai = ti_sai + tp_sai; gr_chua = ti_chua + tp_chua; gr_huy = ti_huy + tp_huy
    gr_ty = f"{round(gr_dung/gr_gd*100, 2):.2f}" if gr_gd else "0.00"
    _append_data_row(t1, grand_total_xml, ['', 'Tổng (I + II):', str(gr_gd), str(gr_dung),
                                           str(gr_sai), str(gr_chua), str(gr_huy), gr_ty])

    t4 = doc.tables[4]
    sec_hdr_xml  = copy.deepcopy(t4.rows[1]._tr)
    data_r4_xml  = copy.deepcopy(t4.rows[2]._tr)
    total_r4_xml = copy.deepcopy(t4.rows[6]._tr)
    grand_r4_xml = copy.deepcopy(t4.rows[13]._tr)

    while len(t4.rows) > 1:
        t4._tbl.remove(t4.rows[-1]._tr)

    def _append_r4(xml_tpl, vals):
        t4._tbl.append(copy.deepcopy(xml_tpl))
        _set_row_cells(t4.rows[-1], vals)

    def _dept_label(name: str) -> str:
        return name.removeprefix('Phòng ').strip()

    _append_r4(sec_hdr_xml, ['I. Hệ thống IPCAS'])
    ti4 = {'gd': 0, 'huy': 0, 'kq': 0, 'cq': 0}
    for dept in dept_data:
        tong = dept.get('ipcas_tong_gd', 0); huy = dept.get('ipcas_bt_huy', 0)
        kq   = dept.get('ipcas_huy_kq', 0);  cq  = dept.get('ipcas_huy_cq', 0)
        pct  = f"{round(cq/tong*100, 2):.2f}" if tong else "0.00"
        ti4['gd'] += tong; ti4['huy'] += huy; ti4['kq'] += kq; ti4['cq'] += cq
        _append_r4(data_r4_xml, [_dept_label(dept.get('dept_name', '')),
                                  str(tong), str(huy), str(kq), str(cq), pct])
    pct_i4 = f"{round(ti4['cq']/ti4['gd']*100, 2):.2f}" if ti4['gd'] else "0.00"
    _append_r4(total_r4_xml, ['Tổng (I)', str(ti4['gd']), str(ti4['huy']), str(ti4['kq']), str(ti4['cq']), pct_i4])

    _append_r4(sec_hdr_xml, ['II. Hệ thống thanh toán tập trung'])
    tp4 = {'gd': 0, 'huy': 0, 'kq': 0, 'cq': 0}
    for dept in dept_data:
        tong = dept.get('payment_tong_gd', 0); huy = dept.get('payment_bt_huy', 0)
        kq   = dept.get('payment_huy_kq', 0);  cq  = dept.get('payment_huy_cq', 0)
        pct  = f"{round(cq/tong*100, 2):.2f}" if tong else "0.00"
        tp4['gd'] += tong; tp4['huy'] += huy; tp4['kq'] += kq; tp4['cq'] += cq
        _append_r4(data_r4_xml, [_dept_label(dept.get('dept_name', '')),
                                  str(tong), str(huy), str(kq), str(cq), pct])
    pct_ii4 = f"{round(tp4['cq']/tp4['gd']*100, 2):.2f}" if tp4['gd'] else "0.00"
    _append_r4(total_r4_xml, ['Tổng (II)', str(tp4['gd']), str(tp4['huy']), str(tp4['kq']), str(tp4['cq']), pct_ii4])

    gr4 = {k: ti4[k] + tp4[k] for k in ti4}
    pct_gr4 = f"{round(gr4['cq']/gr4['gd']*100, 2):.2f}" if gr4['gd'] else "0.00"
    _append_r4(grand_r4_xml, ['Tổng (I + II)', str(gr4['gd']), str(gr4['huy']),
                               str(gr4['kq']), str(gr4['cq']), pct_gr4])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── ZIP nhiều Excel ───────────────────────────────────────────────────────────

def build_zip(files: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)
    buf.seek(0)
    return buf.read()
