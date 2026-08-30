"""Áp quy chuẩn lên file .docx: sửa định dạng, sửa chữ, đánh dấu vùng đã sửa.

## Nguyên tắc: chỉ ghi khi giá trị ĐANG SAI

Không quét một lượt rồi ghi đè tất cả. Lý do là vùng đánh dấu: nếu cứ ghi đè
thì mọi đoạn đều "đã sửa" và cả văn bản vàng khè — người dùng không còn biết
chỗ nào thật sự sai. Nên trước mỗi lần ghi đều so với **giá trị đang có hiệu
lực**, kể cả khi giá trị đó đến từ style chứ không đặt thẳng trên run.

`_hieu_luc_run()` và `_hieu_luc_doan()` làm đúng việc leo ngược chuỗi style để
biết cỡ chữ / căn lề thực sự đang hiển thị. Bỏ bước này mà đọc thẳng
`run.font.size` thì mọi văn bản soạn bằng style (đa số văn bản Word) đều trả
`None`, bị hiểu là "chưa đặt" và bị ghi đè toàn bộ.

## Ba màu đánh dấu

Vàng = sửa định dạng, xanh lá = sửa chữ (viết hoa, đánh số), xanh ngọc = ghép
cụm từ liền dòng. Ba loại này người kiểm tra xử lý khác nhau: sửa định dạng thì
liếc qua là xong, còn sửa chữ thì phải đọc lại xem máy có làm hỏng câu không.
Cùng một màu là bắt họ đọc lại tất cả.
"""
import logging
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph

_log = logging.getLogger(__name__)

_CAN_LE = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_TEN_CAN_LE = {"left": "trái", "center": "giữa", "right": "phải", "justify": "đều hai bên"}

# Sai số bỏ qua khi so số đo — Word lưu lề theo EMU, đổi qua lại mm luôn lệch
# vài phần nghìn. Không có ngưỡng này thì lề 20 mm nào cũng bị coi là sai.
_SAI_SO_MM = 0.3
_SAI_SO_PT = 0.05


def _so(v) -> str:
    """6.0 → "6"; 1.5 → "1,5" — nhật ký cho người đọc, không phải cho máy."""
    f = float(v)
    return (str(int(f)) if f == int(f) else f"{f:g}").replace(".", ",")


# ── Duyệt tài liệu ───────────────────────────────────────────────────────────
def duyet_doan(doc) -> list[tuple[Paragraph, bool]]:
    """Mọi đoạn trong thân văn bản, ĐÚNG thứ tự xuất hiện, kèm cờ nằm-trong-bảng.

    `doc.paragraphs` bỏ qua đoạn nằm trong ô bảng, mà khối Quốc hiệu / tên đơn
    vị đầu trang thường được dựng bằng bảng hai cột. Đi thẳng vào cây XML là
    cách duy nhất giữ được đúng thứ tự giữa đoạn thường và bảng.
    """
    ket_qua: list[tuple[Paragraph, bool]] = []

    # Cha của Paragraph phải là đối tượng CÓ `.part` (Paragraph.style đi qua đó
    # để tra bảng style). Truyền phần tử XML của ô bảng vào sẽ vỡ ở đúng chỗ
    # đọc style — mà chỉ vỡ với văn bản có bảng, nên dễ lọt qua lúc thử.
    def _di(phan_tu, trong_bang: bool):
        for con in phan_tu.iterchildren():
            if con.tag == qn("w:p"):
                ket_qua.append((Paragraph(con, doc), trong_bang))
            elif con.tag == qn("w:tbl"):
                for hang in con.findall(qn("w:tr")):
                    for o in hang.findall(qn("w:tc")):
                        _di(o, True)

    _di(doc.element.body, False)
    return ket_qua


# ── Giá trị đang có hiệu lực (leo ngược chuỗi style) ─────────────────────────
def _hieu_luc_run(run, p: Paragraph, thuoc_tinh: str):
    gia_tri = getattr(run.font, thuoc_tinh)
    if gia_tri is not None:
        return gia_tri
    style = p.style
    for _ in range(10):                       # chặn vòng lặp nếu style tự tham chiếu
        if style is None:
            break
        gia_tri = getattr(style.font, thuoc_tinh, None)
        if gia_tri is not None:
            return gia_tri
        style = style.base_style
    return None


def _hieu_luc_doan(p: Paragraph, thuoc_tinh: str):
    gia_tri = getattr(p.paragraph_format, thuoc_tinh)
    if gia_tri is not None:
        return gia_tri
    style = p.style
    for _ in range(10):
        if style is None:
            break
        gia_tri = getattr(style.paragraph_format, thuoc_tinh, None)
        if gia_tri is not None:
            return gia_tri
        style = style.base_style
    return None


def _dat_phong_chu(run, ten: str) -> None:
    """Đặt phông cho CẢ bốn loại ký tự, không chỉ ascii.

    `run.font.name` của python-docx chỉ ghi w:ascii và w:hAnsi. Chữ tiếng Việt
    có dấu trong file do Word tạo lại thường rơi vào nhánh w:cs (complex
    script); bỏ nhánh đó thì trên máy khác chữ có dấu hiện bằng phông khác chữ
    không dấu — cùng một dòng mà hai kiểu chữ.
    """
    run.font.name = ten
    rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for thuoc in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(thuoc), ten)


# ── Sửa chữ theo từng khoảng, giữ nguyên run ─────────────────────────────────
def ap_sua_text(p: Paragraph, sua: list[tuple[int, int, str]]) -> set[int]:
    """Áp danh sách (đầu, cuối, chữ mới) lên các run của đoạn.

    Trả về chỉ số các run bị đụng tới, để bên gọi tô màu đúng chỗ.

    Đi NGƯỢC từ cuối lên đầu: mỗi lần ghi là độ dài đoạn đổi, mọi chỉ số phía
    sau lệch theo. Đi xuôi thì phải cộng trừ bù, sai một lần là ghi đè lên
    giữa một chữ khác mà không có gì báo.
    """
    if not sua:
        return set()
    runs = p.runs
    if not runs:
        return set()

    # Mốc bắt đầu của từng run trên chuỗi ghép
    moc: list[int] = []
    tong = 0
    for r in runs:
        moc.append(tong)
        tong += len(r.text)

    da_sua: set[int] = set()
    for dau, cuoi, moi in sorted(sua, key=lambda x: x[0], reverse=True):
        if dau >= cuoi or cuoi > tong:
            continue
        lien_quan = [i for i, r in enumerate(runs)
                     if moc[i] < cuoi and moc[i] + len(r.text) > dau]
        if not lien_quan:
            continue
        dau_tien = lien_quan[0]
        r0 = runs[dau_tien]
        cuc_bo_dau = dau - moc[dau_tien]
        if len(lien_quan) == 1:
            cuc_bo_cuoi = cuoi - moc[dau_tien]
            r0.text = r0.text[:cuc_bo_dau] + moi + r0.text[cuc_bo_cuoi:]
        else:
            # Khoảng sửa vắt qua nhiều run: dồn chữ mới vào run ĐẦU (nó mang
            # định dạng của chỗ bắt đầu), xoá phần bị phủ ở các run sau.
            r0.text = r0.text[:cuc_bo_dau] + moi
            for i in lien_quan[1:-1]:
                runs[i].text = ""
            cuoi_cung = lien_quan[-1]
            runs[cuoi_cung].text = runs[cuoi_cung].text[cuoi - moc[cuoi_cung]:]
        da_sua.update(lien_quan)
        # Chuỗi ghép đã đổi → dựng lại mốc cho các lần sửa còn lại (đứng TRƯỚC
        # khoảng vừa ghi nên chỉ cần mốc phía trước, nhưng dựng lại cả cho gọn).
        moc, tong = [], 0
        for r in runs:
            moc.append(tong)
            tong += len(r.text)
    return da_sua


def _to_mau(p: Paragraph, chi_so: set[int] | None, mau: str) -> None:
    try:
        gia_tri = getattr(WD_COLOR_INDEX, mau)
    except AttributeError:
        _log.warning("Màu đánh dấu không hợp lệ: %s — dùng YELLOW", mau)
        gia_tri = WD_COLOR_INDEX.YELLOW
    runs = p.runs
    muc_tieu = range(len(runs)) if chi_so is None else chi_so
    for i in muc_tieu:
        if 0 <= i < len(runs) and runs[i].text:
            runs[i].font.highlight_color = gia_tri


def _xoa_danh_dau(p: Paragraph) -> None:
    for r in p.runs:
        r.font.highlight_color = None


# ── Định dạng một đoạn ───────────────────────────────────────────────────────
def _them(ds: list, loai: str, mo_ta: str) -> None:
    if (loai, mo_ta) not in ds:
        ds.append((loai, mo_ta))


def _dinh_dang_doan(p: Paragraph, ma: str, tp: dict, chung: dict) -> list[tuple[str, str]]:
    """Áp cỡ chữ / kiểu chữ / căn lề cho đoạn.

    Trả `[(loại, mô tả), …]`. `loại` là "chung" khi sửa đổi đó áp đồng loạt
    cho cả văn bản (giãn dòng, cách đoạn, phông chữ…) — những thứ đó vào mục
    "Sửa chung" của nhật ký và KHÔNG bôi màu; "rieng" là khác biệt của riêng
    đoạn đó (cỡ chữ, đậm/nghiêng, căn lề) — có bôi màu.

    Phân loại bằng nhãn trả về chứ không bằng cách đoán tiền tố chuỗi ở bên
    gọi: đổi một chữ trong nhãn là phân loại im lặng sai, không lỗi nào báo.
    """
    ghi_nhan: list[tuple[str, str]] = []
    pf = p.paragraph_format

    # ── Phông chữ và màu chữ: áp cho MỌI đoạn, kể cả ô bảng ──
    if chung.get("ep_phong_chu"):
        ten = chung.get("phong_chu") or "Times New Roman"
        for r in p.runs:
            if r.text and _hieu_luc_run(r, p, "name") != ten:
                _dat_phong_chu(r, ten)
                _them(ghi_nhan, "chung", f"phông chữ → {ten}")
    if chung.get("ep_mau_den"):
        for r in p.runs:
            mau = r.font.color
            if r.text and mau is not None and mau.rgb is not None and str(mau.rgb) != "000000":
                r.font.color.rgb = RGBColor(0, 0, 0)
                _them(ghi_nhan, "chung", "màu chữ → đen")

    # Ô bảng số liệu: cỡ chữ và căn lề do người soạn quyết (Điều 4.2 cho phép
    # bảng biểu trình bày riêng). Dừng ở đây.
    if ma in ("bang", "trong") or not tp:
        return ghi_nhan

    # ── Cỡ chữ ──
    co = tp.get("co")
    if co:
        moi = Pt(float(co))
        for r in p.runs:
            if not r.text:
                continue
            hien = _hieu_luc_run(r, p, "size")
            if hien is None or abs(hien.pt - float(co)) > _SAI_SO_PT:
                r.font.size = moi
                _them(ghi_nhan, "rieng", f"cỡ chữ → {_so(co)}")

    # ── Đậm / nghiêng ──
    for khoa, thuoc, nhan in (("dam", "bold", "chữ đậm"), ("nghieng", "italic", "chữ nghiêng")):
        mong_muon = tp.get(khoa)
        if mong_muon is None:
            continue
        for r in p.runs:
            if not r.text:
                continue
            if bool(_hieu_luc_run(r, p, thuoc)) != bool(mong_muon):
                setattr(r.font, thuoc, bool(mong_muon))
                _them(ghi_nhan, "rieng", f"{nhan} → {'bật' if mong_muon else 'tắt'}")

    # ── Căn lề ──
    can = tp.get("can")
    if can in _CAN_LE:
        hien = _hieu_luc_doan(p, "alignment")
        if hien != _CAN_LE[can]:
            pf.alignment = _CAN_LE[can]
            _them(ghi_nhan, "rieng", f"căn {_TEN_CAN_LE[can]}")

    # ── Thụt dòng đầu và thụt cả đoạn ──
    thut = tp.get("thut_cm")
    if thut is not None:
        hien = _hieu_luc_doan(p, "first_line_indent")
        hien_cm = 0.0 if hien is None else hien.cm
        if abs(hien_cm - float(thut)) > 0.02:
            pf.first_line_indent = Cm(float(thut))
            _them(ghi_nhan, "chung", "thụt dòng đầu theo từng thành phần thể thức")
    le_trai = tp.get("le_trai_cm")
    if le_trai is not None:
        hien = _hieu_luc_doan(p, "left_indent")
        hien_cm = 0.0 if hien is None else hien.cm
        if abs(hien_cm - float(le_trai)) > 0.02:
            pf.left_indent = Cm(float(le_trai))
            _them(ghi_nhan, "chung", "lề trái đoạn theo từng thành phần thể thức")

    # ── Giãn dòng và cách đoạn ──
    # Thành phần khai riêng thì ÉP CHÍNH XÁC (kể cả ép về 0); không khai thì
    # theo giá trị chung và chỉ NÂNG LÊN cho đủ mức tối thiểu Điều 12.6 đòi.
    #
    # Hai cách xử lý khác nhau là có chủ đích. Khối thể thức đầu trang phải ép
    # chính xác về dòng đơn / 0pt — Điều 7.3 và 8.2 nói Quốc hiệu, Tiêu ngữ,
    # tên đơn vị "trình bày cách nhau dòng đơn", nên ở đó "chỉ nâng lên" là sai
    # hướng: file đang giãn 1,5 sẽ được giữ nguyên 1,5 và Tiêu ngữ vẫn nằm xa
    # Quốc hiệu. Ngược lại với lời văn, quy định chỉ nêu mức TỐI THIỂU nên ai
    # để 8pt vẫn hợp lệ, hạ xuống 6pt là sửa thứ không sai.
    gian = tp.get("gian_dong")
    ep_chinh_xac = gian is not None
    if gian is None:
        gian = chung.get("gian_dong")
    if gian:
        hien = _hieu_luc_doan(p, "line_spacing")
        # line_spacing trả float khi là bội số, trả Length khi đặt cứng theo pt
        hien_so = hien if isinstance(hien, (int, float)) else None
        if hien_so is None or abs(float(hien_so) - float(gian)) > 0.01:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = float(gian)
            _them(ghi_nhan, "rieng" if ep_chinh_xac else "chung",
                  f"giãn dòng → {_so(gian)}")

    cach = tp.get("cach_doan_pt")
    ep_chinh_xac = cach is not None
    if cach is None:
        cach = chung.get("cach_doan_pt")
    if cach is not None:
        hien = _hieu_luc_doan(p, "space_after")
        hien_pt = 0.0 if hien is None else hien.pt
        lech = (abs(hien_pt - float(cach)) > _SAI_SO_PT if ep_chinh_xac
                else hien_pt < float(cach) - _SAI_SO_PT)
        if lech:
            pf.space_after = Pt(float(cach))
            _them(ghi_nhan, "rieng" if ep_chinh_xac else "chung",
                  f"cách đoạn → {_so(cach)} pt")
    # ── Khoảng trống TRƯỚC đoạn ──
    # Khoảng cách giữa hai đoạn là `space_after` của đoạn trên CỘNG
    # `space_before` của đoạn dưới. Để cả hai cùng có giá trị thì con số thật
    # không hiện ở đâu cả: file đặt 7pt/7pt cho ra khoảng cách 14pt, mà hộp
    # Paragraph của Word chỉ hiện hai số 7. Đưa `space_before` về 0 để khoảng
    # cách chỉ do `space_after` quyết định — đúng một nguồn, nhìn là biết.
    if chung.get("bo_khoang_truoc_doan"):
        hien_tr = _hieu_luc_doan(p, "space_before")
        if hien_tr is not None and hien_tr.pt > _SAI_SO_PT:
            pf.space_before = Pt(0)
            _them(ghi_nhan, "chung", "bỏ khoảng trống trước đoạn")

    return ghi_nhan


def _ep_hoa_thuong(p: Paragraph, kieu: str | None) -> set[int]:
    """Ép cả đoạn thành in hoa (hoặc in thường). Trả chỉ số run đã đổi.

    Sửa THẮNG từng run chứ không gọi `ap_sua_text()` với một khoảng phủ cả đoạn:
    khoảng phủ nhiều run sẽ dồn chữ về run đầu và làm rỗng các run còn lại — đổi
    hoa/thường vốn không đổi độ dài, không có lý do gì phải trả giá đó.
    """
    if kieu not in ("hoa", "thuong"):
        return set()
    da_doi: set[int] = set()
    for i, r in enumerate(p.runs):
        moi = r.text.upper() if kieu == "hoa" else r.text.lower()
        if moi != r.text:
            r.text = moi
            da_doi.add(i)
    return da_doi


# ── Danh sách tự động của Word ───────────────────────────────────────────────
def _tim_numPr(p: Paragraph):
    """Thẻ w:numPr đang chi phối đoạn — trên chính đoạn, hoặc trên style của nó.

    Bấm nút bullet trên thanh công cụ Word thì thẻ nằm trên ĐOẠN. Nhưng chọn
    style "List Bullet" / "List Paragraph" (và mọi đoạn do python-docx tạo với
    `style=`) thì thẻ nằm trên STYLE, đoạn không có gì cả. Chỉ đọc trên đoạn là
    bỏ sót đúng nửa số văn bản — mà bỏ sót thì không lỗi, chỉ là dấu chấm tròn
    nằm nguyên đó và không ai hiểu vì sao phần mềm "không làm gì".
    """
    pPr = p._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return pPr.numPr
    style = p.style
    for _ in range(10):
        if style is None:
            break
        el = getattr(style, "element", None)
        st_pPr = el.find(qn("w:pPr")) if el is not None else None
        if st_pPr is not None:
            numPr = st_pPr.find(qn("w:numPr"))
            if numPr is not None:
                return numPr
        style = style.base_style
    return None


def _kieu_danh_so(doc, p: Paragraph) -> str | None:
    """'bullet' / 'so' / None — đoạn có dùng đánh số tự động của Word không.

    Số hiển thị trên màn hình do Word tự tính, KHÔNG nằm trong file dưới dạng
    chữ. Đọc bằng python-docx chỉ thấy một tham chiếu numId. Nên chỉ chuyển
    được loại 'bullet' (dấu chấm tròn — không có số nào phải tính); loại có số
    thì đổi thành chữ gõ tay đồng nghĩa với tự đếm lại toàn bộ, sai một chỗ là
    lệch số cả văn bản mà không ai biết. Mặc định để nguyên và ghi cảnh báo.
    """
    numPr = _tim_numPr(p)
    if numPr is None:
        return None
    nut_id = numPr.find(qn("w:numId"))
    if nut_id is None:
        return None
    num_id = nut_id.get(qn("w:val"))
    nut_lvl = numPr.find(qn("w:ilvl"))
    muc = nut_lvl.get(qn("w:val")) if nut_lvl is not None else "0"
    if num_id in (None, "0"):
        return None
    try:
        goc = doc.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError, ValueError):
        return "khong_ro"

    abs_id = None
    for num in goc.findall(qn("w:num")):
        if num.get(qn("w:numId")) == num_id:
            el = num.find(qn("w:abstractNumId"))
            abs_id = el.get(qn("w:val")) if el is not None else None
            break
    if abs_id is None:
        return "khong_ro"

    for abs_num in goc.findall(qn("w:abstractNum")):
        if abs_num.get(qn("w:abstractNumId")) != abs_id:
            continue
        for lvl in abs_num.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != muc:
                continue
            fmt = lvl.find(qn("w:numFmt"))
            val = fmt.get(qn("w:val")) if fmt is not None else None
            return "bullet" if val == "bullet" else "so"
    return "khong_ro"


def _go_danh_so_tu_dong(doc, p: Paragraph) -> None:
    """Bỏ đánh số tự động của đoạn, giữ nguyên chữ.

    Hai đường khác nhau tuỳ thẻ w:numPr nằm ở đâu:

    * Trên đoạn → xoá thẳng thẻ đó.
    * Trên STYLE → KHÔNG được xoá, vì style dùng chung: xoá một lần là mọi đoạn
      khác cùng style cũng mất dấu đầu dòng, kể cả đoạn phần mềm chưa xét tới.
      Thay vào đó đổi đoạn này về style thường — cách đó cũng gỡ luôn phần thụt
      lề riêng của style danh sách, thứ mà quy định không cho phép.
    """
    pPr = p._p.pPr
    if pPr is not None and pPr.numPr is not None:
        pPr.remove(pPr.numPr)
        return
    try:
        p.style = doc.styles["Normal"]
    except KeyError:                       # tài liệu không có style "Normal"
        _log.warning("Không đổi được style danh sách về Normal")


# ── Số trang ─────────────────────────────────────────────────────────────────
def _them_so_trang(section, co_chu: float, phong: str) -> bool:
    """Chèn số trang canh giữa vào lề trên, bỏ trang đầu (Điều 4.4).

    Chỉ chèn khi header đang TRỐNG. Header có sẵn thường là logo hoặc dòng chỉ
    dẫn của đơn vị — ghi đè lên đó là xoá mất nội dung người dùng cố ý đặt.
    """
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    if any(dp.text.strip() for dp in header.paragraphs):
        return False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    if p.runs:
        return False
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for kieu, noi_dung in (("begin", None), (None, " PAGE "), ("end", None)):
        r = p.add_run()
        _dat_phong_chu(r, phong)
        r.font.size = Pt(float(co_chu))
        if kieu:
            r._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="{kieu}"/>'))
        else:
            r._r.append(parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve">{noi_dung}</w:instrText>'))
    return True


def dat_trang(doc, cfg_trang: dict) -> list[str]:
    """Khổ giấy, định lề, số trang — áp cho MỌI section. Trả mô tả đã sửa."""
    if not cfg_trang.get("ap_dung"):
        return []
    ghi_nhan: list[str] = []
    do = [
        ("page_width", "rong_mm", "khổ giấy rộng"),
        ("page_height", "cao_mm", "khổ giấy cao"),
        ("top_margin", "le_tren_mm", "lề trên"),
        ("bottom_margin", "le_duoi_mm", "lề dưới"),
        ("left_margin", "le_trai_mm", "lề trái"),
        ("right_margin", "le_phai_mm", "lề phải"),
    ]
    for section in doc.sections:
        for thuoc, khoa, nhan in do:
            mong = cfg_trang.get(khoa)
            if mong is None:
                continue
            hien = getattr(section, thuoc)
            if hien is None or abs(hien.mm - float(mong)) > _SAI_SO_MM:
                setattr(section, thuoc, Mm(float(mong)))
                mo_ta = f"{nhan} → {_so(mong)} mm"
                if mo_ta not in ghi_nhan:
                    ghi_nhan.append(mo_ta)
        if cfg_trang.get("danh_so_trang"):
            try:
                if _them_so_trang(section, cfg_trang.get("co_so_trang") or 14,
                                  "Times New Roman"):
                    if "đánh số trang (bỏ trang đầu)" not in ghi_nhan:
                        ghi_nhan.append("đánh số trang (bỏ trang đầu)")
            except Exception as e:                        # noqa: BLE001
                _log.warning("Không chèn được số trang: %s", e)
    return ghi_nhan
