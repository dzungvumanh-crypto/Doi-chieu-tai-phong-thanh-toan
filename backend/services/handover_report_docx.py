"""Xuất báo cáo bàn giao chứng từ ra Word (A4 ngang).

Nhận nguyên dict của `handover_report_service.compute_period()` — không tự
truy vấn DB, không tự tính lại — để file Word và màn hình luôn cùng một con số.
"""
import io
from datetime import date
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
TITLE_SIZE = Pt(16)

# Cột bảng tổng hợp / bảng chi tiết (cm) — vừa khổ A4 ngang trừ lề
_SUMMARY_WIDTHS = [1.5, 9.0, 3.5, 3.5, 3.5, 3.5]
_DETAIL_WIDTHS = [1.5, 7.0, 3.5, 3.5, 3.5, 2.5]

# Cột "Họ và tên" ở bảng chi tiết — gộp ô theo cán bộ
_NAME_COL = 1


def _fmt_date(iso: Optional[str]) -> str:
    """2026-07-02 → 02/07/2026"""
    if not iso:
        return ""
    try:
        return date.fromisoformat(str(iso)[:10]).strftime("%d/%m/%Y")
    except ValueError:
        return str(iso)


def _fmt_rate(rate) -> str:
    return f"{rate:.1f}%" if rate is not None else "—"


def _setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    # Đổi orientation KHÔNG tự hoán đổi kích thước trang — phải gán tay,
    # nếu không Word vẫn in ra khổ dọc.
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    # Font chữ Việt chỉ ăn khi set cả rFonts eastAsia/cs, không chỉ style.font.name
    rpr = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:eastAsia", "w:cs", "w:ascii", "w:hAnsi"):
        rpr.set(qn(attr), FONT_NAME)


def _add_title(doc: Document, year: int, month: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"BÁO CÁO BÀN GIAO CHỨNG TỪ THÁNG {month:02d} NĂM {year:04d}")
    run.bold = True
    run.font.size = TITLE_SIZE
    p.paragraph_format.space_after = Pt(14)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _make_table(doc: Document, headers: list, widths: list):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, text, w in zip(table.rows[0].cells, headers, widths):
        cell.width = Cm(w)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
    return table


def _fill_row(table, values: list, widths: list, bold: bool = False) -> None:
    """Thêm một dòng; cột đầu tiên (STT) và các cột số căn giữa, cột chữ căn trái."""
    cells = table.add_row().cells
    for idx, (cell, val, w) in enumerate(zip(cells, values, widths)):
        cell.width = Cm(w)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(val))
        run.bold = bold


def _add_summary(doc: Document, data: dict) -> None:
    overall = data.get("overall", {}) or {}
    _add_heading(doc, "I. TỔNG HỢP")

    table = _make_table(
        doc,
        ["STT", "Phòng nghiệp vụ", "Tổng chứng từ", "Nộp đúng hạn", "Nộp quá hạn", "Tỷ lệ đúng hạn"],
        _SUMMARY_WIDTHS,
    )
    for stt, row in enumerate(data.get("by_dept", []) or [], 1):
        _fill_row(table, [
            stt,
            row.get("dept_name", ""),
            row.get("total", 0),
            row.get("on_time", 0),
            row.get("late", 0),
            _fmt_rate(row.get("rate")),
        ], _SUMMARY_WIDTHS)

    _fill_row(table, [
        "",
        "TỔNG CỘNG",
        overall.get("total", 0),
        overall.get("on_time", 0),
        overall.get("late", 0),
        _fmt_rate(overall.get("rate")),
    ], _SUMMARY_WIDTHS, bold=True)


def _set_cell_text(cell, text: str, align, bold: bool = False) -> None:
    """Ghi đè nội dung một ô — dùng sau khi merge.

    `merge()` nối paragraph của MỌI ô bị gộp lại, nên ô gộp 7 dòng có 7 paragraph
    rỗng và cao gấp 7 lần. Phải dọn hết rồi mới ghi.
    """
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    para = cell.paragraphs[0]
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    para.alignment = align
    r = para.add_run(text)
    r.bold = bold


def _group_by_staff(rows: list) -> list:
    """Gom chứng từ của cùng một cán bộ vào một cụm liên tiếp.

    Khoá gom là `staff_id` chứ không phải tên: hai cán bộ trùng họ tên sẽ bị gộp
    nhầm thành một ô. Thứ tự cụm giữ theo lần xuất hiện đầu tiên (backend đã sort
    theo số ngày chậm giảm dần), trong cụm sắp theo ngày giao dịch.
    """
    groups: dict = {}
    for e in rows:
        key = e.get("staff_id") or f"name:{e.get('staff_name', '')}"
        groups.setdefault(key, []).append(e)
    return [sorted(g, key=lambda e: e.get("transaction_date") or "") for g in groups.values()]


def _add_late_detail(doc: Document, data: dict) -> None:
    _add_heading(doc, "II. CHI TIẾT CHỨNG TỪ NỘP QUÁ HẠN THEO PHÒNG")

    late_entries = data.get("late_entries", []) or []
    if not late_entries:
        doc.add_paragraph("Không có chứng từ nào nộp quá hạn trong kỳ.")
        return

    by_dept: dict = {}
    for e in late_entries:
        by_dept.setdefault(e.get("dept_name", ""), []).append(e)

    headers = ["STT", "Họ và tên", "Ngày giao dịch", "Ngày nộp", "Số ngày chậm", "Số tờ"]
    for idx, (dept_name, rows) in enumerate(by_dept.items(), 1):
        _add_heading(doc, f"{idx}. {dept_name} — {len(rows)} chứng từ quá hạn")
        table = _make_table(doc, headers, _DETAIL_WIDTHS)

        stt = 1
        for group in _group_by_staff(rows):
            first_row = len(table.rows)
            for e in group:
                # Cột họ tên để trống, điền sau khi merge
                _fill_row(table, [
                    stt, "",
                    _fmt_date(e.get("transaction_date")),
                    _fmt_date(e.get("submitted_date")),
                    e.get("days_late", 0),
                    e.get("sheet_count", 0),
                ], _DETAIL_WIDTHS)
                stt += 1

            last_row = len(table.rows) - 1
            cell = table.cell(first_row, _NAME_COL)
            if last_row > first_row:
                cell = cell.merge(table.cell(last_row, _NAME_COL))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_text(cell, group[0].get("staff_name", ""), WD_ALIGN_PARAGRAPH.LEFT)


def build_report_docx(data: dict, year: int, month: int) -> bytes:
    doc = Document()
    _setup_page(doc)
    _add_title(doc, year, month)
    _add_summary(doc, data)
    _add_late_detail(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
