"""Tách 18 mẫu trình bày văn bản từ Phụ lục V ra 18 file .docx riêng.

Chạy tay mỗi khi Phụ lục V có bản mới; kết quả nằm ở `templates/vb_mau/` và
được commit vào repo — máy chính KHÔNG chạy script này.

    python scripts/tach_mau_vb.py

## Vì sao xoá bớt trên bản gốc chứ không chép sang file trắng

Chép từng đoạn sang một `Document()` mới thì mất `styles.xml`, `numbering.xml`
và `theme1.xml` — đoạn chép sang rơi về style Normal của Word, tức là mất luôn
cỡ chữ và giãn dòng mà chính cái mẫu đang minh hoạ. Mở bản gốc rồi xoá những
phần tử ngoài phạm vi giữ được nguyên cả ba phần đó, kể cả footnote.

## Vì sao gỡ header/footer

Sổ `sectPr` của Phụ lục trỏ tới `header3.xml` — trong đó có trường PAGE và kết
quả cũ "22", tức số trang của Phụ lục. Giữ lại thì mẫu trắng nào cũng đội một
số trang lạ trên đầu. Nặng hơn: `_them_so_trang()` bên `ap_dung.py` chỉ đánh số
khi header còn TRỐNG, nên văn bản soạn từ mẫu này sẽ không bao giờ được đánh
số trang — hỏng lặng lẽ, không lỗi, không log.

## Vì sao ép lại lề trang

File Phụ lục để lề trên/dưới 15 mm cho mỗi mẫu vừa gọn một trang giấy. Quy
chuẩn của dự án là 20 mm (`quy_chuan.py`, theo Điều 4). Giữ nguyên lề của Phụ
lục là phát ra mẫu trắng sai lề ngay từ dòng đầu.
"""
import json
import os
import re
import sys
import unicodedata

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document                                          # noqa: E402
from docx.shared import Mm                                         # noqa: E402

from backend.core.paths import PROJECT_ROOT, resolve_path          # noqa: E402
from backend.services.vb_format.quy_chuan import mac_dinh          # noqa: E402

NGUON = resolve_path(PROJECT_ROOT, "979-QyD-NHNo-PC (Trình bày VB)",
                     "Phụ lục V_Mẫu trình bày văn bản.docx")
DICH = os.path.join(PROJECT_ROOT, "templates", "vb_mau")

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_HEADING = re.compile(r"^M[ẫâa]u\s*(\d+)\s*:", re.IGNORECASE)

# Đoạn "trống" theo nghĩa xoá được: không chữ VÀ không chứa đối tượng vẽ.
# Đường kẻ ngang dưới Tiêu ngữ là <v:line> nằm trong một đoạn KHÔNG có chữ —
# lọc bằng mỗi `.text` là xoá mất đường kẻ mà không ai thấy cho tới lúc in.
_CO_HINH = ("v:line", "w:drawing", "w:pict", "w:object")


# ── Tên file thuần ASCII ─────────────────────────────────────────────────────
def _slug(s: str, gioi_han: int = 40) -> str:
    s = s.replace("đ", "d").replace("Đ", "D")
    s = "".join(c for c in unicodedata.normalize("NFD", s)
                if not unicodedata.combining(c))
    s = re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_").lower()
    return s[:gioi_han].rstrip("_")


# ── Đọc bảng mục lục ở trang đầu ─────────────────────────────────────────────
def _muc_luc(doc) -> dict[int, str]:
    """{1: 'Nghị quyết', …} lấy từ bảng MẪU SỐ / TÊN MẪU."""
    ten: dict[int, str] = {}
    for row in doc.tables[0].rows:
        o = [c.text.strip() for c in row.cells]
        m = re.match(r"^M[ẫâa]u\s*(\d+)$", o[0], re.IGNORECASE) if len(o) >= 2 else None
        if m:
            ten[int(m.group(1))] = " ".join(o[1].split())
    return ten


def _la_doan_trong(el) -> bool:
    if el.tag != _W + "p":
        return False
    if "".join(el.itertext()).strip():
        return False
    return not any(t in el.xml for t in _CO_HINH)


def _don_sect_pr(doc, cfg_trang: dict) -> None:
    """Gỡ mọi header/footer và ép lề về quy chuẩn, cho mọi section còn lại."""
    for sect in doc.sections:
        for ref in list(sect._sectPr):
            if ref.tag in (_W + "headerReference", _W + "footerReference"):
                sect._sectPr.remove(ref)
        sect.top_margin = Mm(float(cfg_trang["le_tren_mm"]))
        sect.bottom_margin = Mm(float(cfg_trang["le_duoi_mm"]))
        sect.left_margin = Mm(float(cfg_trang["le_trai_mm"]))
        sect.right_margin = Mm(float(cfg_trang["le_phai_mm"]))


def _tach_mot(so: int, ten: str, cfg_trang: dict) -> dict:
    doc = Document(NGUON)
    body = doc.element.body
    kids = list(body.iterchildren())

    moc = [i for i, k in enumerate(kids)
           if k.tag == _W + "p" and _HEADING.match("".join(k.itertext()).strip())]
    thu_tu = [int(_HEADING.match("".join(kids[i].itertext()).strip()).group(1)) for i in moc]
    vi_tri = thu_tu.index(so)
    dau = moc[vi_tri]
    cuoi = moc[vi_tri + 1] if vi_tri + 1 < len(moc) else len(kids)

    # Giữ phạm vi [dau+1, cuoi) — bỏ chính dòng "Mẫu NN:". Dấu ngắt trang nằm
    # trong dòng đó nên xoá dòng là xoá luôn trang trắng đứng trước.
    giu = kids[dau + 1:cuoi]
    while giu and _la_doan_trong(giu[0]):
        giu.pop(0)
    while giu and _la_doan_trong(giu[-1]):
        giu.pop()

    giu_id = {id(k) for k in giu}
    for k in kids:
        if id(k) not in giu_id and k.tag != _W + "sectPr":
            body.remove(k)

    _don_sect_pr(doc, cfg_trang)

    ten_file = f"{so:02d}_{_slug(ten)}.docx"
    doc.save(os.path.join(DICH, ten_file))
    return {"so": so, "ten": ten, "file": ten_file, "phan_tu": len(giu)}


def main() -> int:
    # Bảng điều khiển Windows mặc định cp1252, in tên mẫu tiếng Việt là ném
    # UnicodeEncodeError sau khi ĐÃ ghi xong file — thất bại nhìn như lỗi tách.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, OSError):
        pass

    if not os.path.exists(NGUON):
        print(f"Không thấy file nguồn: {NGUON}")
        return 1
    os.makedirs(DICH, exist_ok=True)

    cfg_trang = mac_dinh()["trang"]
    ten = _muc_luc(Document(NGUON))
    if len(ten) != 18:
        print(f"Bảng mục lục đọc ra {len(ten)} dòng, chờ 18 — dừng.")
        return 1

    muc_luc = [_tach_mot(so, ten[so], cfg_trang) for so in sorted(ten)]
    for m in muc_luc:
        print(f"  {m['so']:02d}  {m['ten'][:48]:<50} {m['file']}  ({m['phan_tu']} phần tử)")

    with open(os.path.join(DICH, "muc_luc.json"), "w", encoding="utf-8") as f:
        json.dump([{k: m[k] for k in ("so", "ten", "file")} for m in muc_luc],
                  f, ensure_ascii=False, indent=2)
    print(f"\nĐã ghi {len(muc_luc)} mẫu vào {DICH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
