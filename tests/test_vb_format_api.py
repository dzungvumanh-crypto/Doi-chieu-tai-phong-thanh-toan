"""API Chuẩn hoá văn bản — quyền, luồng chuẩn hoá → tải về, lưu cấu hình."""
import io
import sqlite3

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.core.deps import get_current_staff
from backend.core.enums import StaffRole
from backend.database import get_db
from backend.main import app

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _file_word() -> bytes:
    doc = Document()
    doc.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    doc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    doc.add_paragraph("1) Quy định này áp dụng cho toàn hệ thống.")
    ra = io.BytesIO()
    doc.save(ra)
    return ra.getvalue()


@pytest.fixture
def client(tmp_path, monkeypatch):
    """TestClient với DB trong RAM đã có bảng vb_format_config và user_tttt.

    Không dùng `admin_client` chung: endpoint cấu hình JOIN sang `user_tttt` để
    hiện "ai sửa lần cuối", mà DB rỗng của fixture chung không có bảng nào.

    `TEMP_DIR` bị trỏ sang thư mục tạm của pytest: endpoint ghi file kết quả ra
    đĩa thật, để nguyên thì mỗi lần chạy test lại vứt thêm rác vào `data/` của
    dự án — và rác đó chỉ bị dọn lúc 23h.
    """
    from backend.api import vb_format as vb_api
    monkeypatch.setattr(vb_api, "TEMP_DIR", tmp_path / "temp_vb_format")
    # check_same_thread=False: FastAPI chạy endpoint `def` trong threadpool, kết
    # nối được tạo ở thread của test nhưng dùng ở thread khác.
    conn = sqlite3.connect(":memory:", check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.executescript(
        """
        CREATE TABLE user_tttt (id INTEGER PRIMARY KEY, full_name TEXT);
        INSERT INTO user_tttt (id, full_name) VALUES (1, 'Test Admin');
        CREATE TABLE vb_format_config (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            config_json TEXT NOT NULL DEFAULT '{}',
            updated_at DATETIME,
            updated_by INTEGER REFERENCES user_tttt(id)
        );
        INSERT INTO vb_format_config (id, config_json) VALUES (1, '{}');
        CREATE TABLE audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT, actor_id INTEGER, action TEXT,
            target_type TEXT, target_id INTEGER, detail TEXT, ip_address TEXT,
            created_at DATETIME
        );
        """
    )
    app.dependency_overrides[get_current_staff] = lambda: {
        "id": 1, "role": StaffRole.ADMIN, "username": "admin", "full_name": "Test Admin"}
    app.dependency_overrides[get_db] = lambda: conn
    yield TestClient(app)
    app.dependency_overrides.clear()
    conn.close()


def test_lay_cau_hinh_tra_du_thu_man_hinh_can(client):
    r = client.get("/api/vb-format/cau-hinh")
    assert r.status_code == 200
    d = r.json()
    assert d["cau_hinh"]["chung"]["phong_chu"] == "Times New Roman"
    assert len(d["nhan"]) == len(d["cau_hinh"]["thanh_phan"]), \
        "mỗi thành phần thể thức phải có nhãn tiếng Việt, nếu không màn cấu hình bỏ sót"
    assert d["dai_co_chu"]["noi_nhan_ds"] == [11, 11]
    assert "YELLOW" in d["mau_danh_dau"]


def test_luu_cau_hinh_chi_ghi_phan_khac_mac_dinh(client):
    r = client.put("/api/vb-format/cau-hinh",
                   json={"thanh_phan": {"noi_dung": {"co": 13}}})
    assert r.status_code == 200
    assert r.json()["cau_hinh"]["thanh_phan"]["noi_dung"]["co"] == 13

    luu = client.app.dependency_overrides[get_db]().execute(
        "SELECT config_json FROM vb_format_config WHERE id = 1").fetchone()[0]
    assert "phong_chu" not in luu, "giá trị trùng mặc định không được lưu lại"
    assert "13" in luu


def test_khoi_phuc_mac_dinh(client):
    client.put("/api/vb-format/cau-hinh", json={"thanh_phan": {"noi_dung": {"co": 13}}})
    r = client.post("/api/vb-format/cau-hinh/mac-dinh")
    assert r.status_code == 200
    assert r.json()["cau_hinh"]["thanh_phan"]["noi_dung"]["co"] == 14


def test_chuan_hoa_roi_tai_ve(client):
    r = client.post("/api/vb-format/chuan-hoa",
                    files={"file": ("quyet_dinh.docx", _file_word(), _DOCX_MIME)})
    assert r.status_code == 200, r.text
    d = r.json()
    assert d["ten_file"] == "quyet_dinh_da_chuan_hoa.docx"
    assert d["thong_ke"]["tong_doan"] == 4

    tai = client.get(f"/api/vb-format/tai-ve/{d['token']}")
    assert tai.status_code == 200
    assert "quyet_dinh_da_chuan_hoa.docx" in tai.headers["content-disposition"]
    doc = Document(io.BytesIO(tai.content))
    assert any(p.text.strip().startswith("1. Quy định này") for p in doc.paragraphs)


def test_cau_hinh_da_luu_duoc_ap_khi_chuan_hoa(client):
    client.put("/api/vb-format/cau-hinh", json={"thanh_phan": {"khoan": {"co": 13}}})
    r = client.post("/api/vb-format/chuan-hoa",
                    files={"file": ("a.docx", _file_word(), _DOCX_MIME)})
    tai = client.get(f"/api/vb-format/tai-ve/{r.json()['token']}")
    doc = Document(io.BytesIO(tai.content))
    khoan = next(p for p in doc.paragraphs if p.text.strip().startswith("1."))
    assert khoan.runs[0].font.size.pt == 13


def test_tu_choi_file_khong_phai_docx(client):
    r = client.post("/api/vb-format/chuan-hoa",
                    files={"file": ("bao_cao.doc", b"khong phai docx", "application/msword")})
    assert r.status_code == 400
    assert ".docx" in r.json()["detail"]


def test_file_hong_bao_loi_de_hieu_khong_tra_500(client):
    r = client.post("/api/vb-format/chuan-hoa",
                    files={"file": ("hong.docx", b"day khong phai zip", _DOCX_MIME)})
    assert r.status_code == 400
    assert "Word" in r.json()["detail"]


def test_token_khong_ton_tai_tra_404(client):
    assert client.get("/api/vb-format/tai-ve/khongcothat").status_code == 404


def test_token_khong_thoat_ra_ngoai_thu_muc(client):
    """Token đi thẳng vào đường dẫn — không được leo ra thư mục cha."""
    r = client.get("/api/vb-format/tai-ve/..%2F..%2Fbackend%2Fmain.py")
    assert r.status_code in (307, 404)


@pytest.fixture
def client_chi_xem(tmp_path, monkeypatch):
    """TestClient của một chuyên viên CHỈ được cấp `menu.vb_format`.

    Cố ý KHÔNG override `require_feature` — để nó chạy thật trên bảng
    `group_features`. Bản test trước override đúng cái dependency đang cần
    kiểm tra rồi khẳng định "nó trả 403": tức là chỉ chứng minh cái giả lập
    của chính mình chạy đúng. Gắn nhầm mã feature vào route thì test đó vẫn
    xanh — mà gắn nhầm mã chính là lỗi cần bắt.

    Phải trỏ `TEMP_DIR` sang thư mục tạm của pytest giống fixture `client`:
    test dưới đây có gọi `/chuan-hoa`, mà endpoint đó ghi file kết quả ra đĩa
    thật — quên patch là mỗi lần chạy test lại vứt thêm rác vào `data/`.
    """
    from backend.api import vb_format as vb_api
    monkeypatch.setattr(vb_api, "TEMP_DIR", tmp_path / "temp_vb_format")
    conn = sqlite3.connect(":memory:", check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.executescript(
        """
        CREATE TABLE user_tttt (id INTEGER PRIMARY KEY, full_name TEXT, role TEXT);
        INSERT INTO user_tttt VALUES (2, 'Chuyên viên', 'chuyen_vien');
        CREATE TABLE user_groups (id INTEGER PRIMARY KEY, name TEXT, is_active INTEGER DEFAULT 1);
        INSERT INTO user_groups VALUES (1, 'Văn thư', 1);
        CREATE TABLE group_members (group_id INTEGER, staff_id INTEGER);
        INSERT INTO group_members VALUES (1, 2);
        CREATE TABLE group_features (group_id INTEGER, feature_code TEXT);
        INSERT INTO group_features VALUES (1, 'menu.vb_format');
        CREATE TABLE vb_format_config (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            config_json TEXT NOT NULL DEFAULT '{}',
            updated_at DATETIME, updated_by INTEGER);
        INSERT INTO vb_format_config (id, config_json) VALUES (1, '{}');
        CREATE TABLE audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT, actor_id INTEGER, action TEXT,
            target_type TEXT, target_id INTEGER, detail TEXT, ip_address TEXT,
            created_at DATETIME);
        CREATE TABLE login_sessions (staff_id INTEGER, ip_address TEXT);
        """
    )
    app.dependency_overrides[get_current_staff] = lambda: {
        "id": 2, "role": StaffRole.CHUYEN_VIEN, "username": "cv", "full_name": "Chuyên viên"}
    app.dependency_overrides[get_db] = lambda: conn
    yield TestClient(app), conn
    app.dependency_overrides.clear()
    conn.close()


def test_chi_co_quyen_menu_thi_khong_ghi_duoc_tham_so(client_chi_xem):
    """Người chỉ có `menu.vb_format` phải bị chặn ở CẢ HAI đường ghi."""
    c, conn = client_chi_xem
    assert c.put("/api/vb-format/cau-hinh",
                 json={"chung": {"gian_dong": 1.0}}).status_code == 403
    assert c.post("/api/vb-format/cau-hinh/mac-dinh").status_code == 403
    # Khẳng định trên DỮ LIỆU, không chỉ trên mã trạng thái: 403 mà vẫn ghi
    # được thì mã trạng thái chẳng nói lên điều gì.
    assert conn.execute(
        "SELECT config_json FROM vb_format_config WHERE id = 1").fetchone()[0] == "{}"


def test_chi_co_quyen_menu_van_xem_va_chuan_hoa_duoc(client_chi_xem):
    """Chặn ghi nhưng KHÔNG được chặn việc dùng — họ vẫn phải chuẩn hoá được."""
    c, _ = client_chi_xem
    assert c.get("/api/vb-format/cau-hinh").status_code == 200
    r = c.post("/api/vb-format/chuan-hoa",
               files={"file": ("a.docx", _file_word(), _DOCX_MIME)})
    assert r.status_code == 200
