"""18 mẫu trình bày sẵn tách từ Phụ lục V — kiểm bản đã commit trong repo.

Đây là dữ liệu tĩnh do `scripts/tach_mau_vb.py` sinh ra, không phải kết quả
tính lúc chạy. Test canh đúng những chỗ hỏng mà mở file ra nhìn KHÔNG thấy
ngay: lề trang sai, còn sót header số trang của Phụ lục, mất footnote.
"""
import json
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document

from backend.services.vb_format.quy_chuan import mac_dinh

MAU_DIR = Path(__file__).resolve().parents[1] / "templates" / "vb_mau"
_HEADING = re.compile(r"^\s*M[ẫâa]u\s*\d+\s*:", re.IGNORECASE | re.MULTILINE)


@pytest.fixture(scope="module")
def muc_luc() -> list[dict]:
    return json.loads((MAU_DIR / "muc_luc.json").read_text(encoding="utf-8"))


def _toan_van(doc) -> str:
    phan = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        phan += [c.text for r in t.rows for c in r.cells]
    return "\n".join(phan)


def test_du_18_mau_va_du_file(muc_luc):
    assert [m["so"] for m in muc_luc] == list(range(1, 19))
    assert all((MAU_DIR / m["file"]).exists() for m in muc_luc)
    # Tên file phải thuần ASCII — xem docs/DESIGN.md, mục đường dẫn NFC/NFD.
    assert all(m["file"].isascii() for m in muc_luc)


def test_le_trang_dung_quy_chuan(muc_luc):
    """Phụ lục V để lề trên/dưới 15 mm cho vừa trang; quy chuẩn là 20 mm."""
    cfg = mac_dinh()["trang"]
    for m in muc_luc:
        s = Document(MAU_DIR / m["file"]).sections[0]
        do = [(s.top_margin, "le_tren_mm"), (s.bottom_margin, "le_duoi_mm"),
              (s.left_margin, "le_trai_mm"), (s.right_margin, "le_phai_mm")]
        for hien, khoa in do:
            assert abs(hien.mm - cfg[khoa]) < 0.5, f"{m['file']} sai {khoa}"


def test_khong_con_header_footer_cua_phu_luc(muc_luc):
    """Header của Phụ lục có trường PAGE. Giữ lại thì `_them_so_trang()` bên
    `ap_dung.py` thấy header đã có chữ và BỎ QUA việc đánh số trang."""
    for m in muc_luc:
        xml = zipfile.ZipFile(MAU_DIR / m["file"]).read("word/document.xml").decode("utf-8")
        assert "headerReference" not in xml, m["file"]
        assert "footerReference" not in xml, m["file"]


def test_da_bo_dong_tieu_de_mau(muc_luc):
    for m in muc_luc:
        assert not _HEADING.search(_toan_van(Document(MAU_DIR / m["file"]))), m["file"]


def test_giu_nguyen_ghi_chu_va_footnote(muc_luc):
    """Người dùng chốt: KHÔNG bỏ phần ghi chú của Phụ lục."""
    ten = {m["so"]: MAU_DIR / m["file"] for m in muc_luc}

    # Ghi chú dạng đoạn văn — Mẫu 07 liệt kê các loại văn bản dùng chung mẫu.
    assert "Ghi chú:" in _toan_van(Document(ten[7]))

    # Ghi chú dạng footnote — Mẫu 04 có 3, Mẫu 17 có 4.
    for so, cho in ((4, 3), (17, 4)):
        xml = zipfile.ZipFile(ten[so]).read("word/document.xml").decode("utf-8")
        assert xml.count("footnoteReference") == cho, f"Mẫu {so}"


def test_moi_mau_con_khoi_the_thuc(muc_luc):
    """Cắt sai phạm vi thì file vẫn mở được nhưng rỗng ruột — không lỗi nào."""
    for m in muc_luc:
        txt = _toan_van(Document(MAU_DIR / m["file"]))
        # Mẫu 18 (Phụ lục văn bản) ngắn nhất — 144 ký tự, chỉ có tiêu đề và
        # dòng "Kèm theo". Ngưỡng lấy dưới nó một chút, đủ bắt trường hợp cắt
        # hụt còn lại vài dòng lẻ.
        assert len(txt.strip()) > 120, f"{m['file']} quá ngắn, nghi cắt hụt"
        assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in txt or "Phụ lục" in txt, m["file"]


# ── API ──────────────────────────────────────────────────────────────────────
def test_api_liet_ke_va_tai_mau(admin_client, muc_luc):
    r = admin_client.get("/api/vb-format/mau")
    assert r.status_code == 200
    assert [m["so"] for m in r.json()] == [m["so"] for m in muc_luc]

    r = admin_client.get("/api/vb-format/mau/8")
    assert r.status_code == 200
    assert r.content[:2] == b"PK"                      # .docx là file zip
    assert "Mau 08" in r.headers["content-disposition"]
    assert len(r.content) == (MAU_DIR / "08_cong_van.docx").stat().st_size


def test_api_mau_khong_ton_tai_tra_404(admin_client):
    assert admin_client.get("/api/vb-format/mau/99").status_code == 404
    # Chuỗi lạ không lọt vào đường dẫn file: kiểu `int` của tham số chặn trước
    # khi handler chạy (422), còn chuỗi có dấu "/" thì không khớp route (404).
    # Tên file luôn lấy từ mục lục, không bao giờ ghép từ chuỗi người dùng gửi.
    assert admin_client.get("/api/vb-format/mau/abc").status_code == 422
    assert admin_client.get(
        "/api/vb-format/mau/..%2F..%2Fsecret").status_code in (404, 422)
