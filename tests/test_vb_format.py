"""Chuẩn hoá văn bản theo QĐ 979 — nhận diện thể thức, sửa chữ, đánh dấu.

Văn bản mẫu trong `_van_ban_sai()` được dựng cố ý sai đủ kiểu hay gặp: phông
Arial, cỡ 11, lề mặc định của Word, gạch đầu dòng bằng "•", khoản đánh "1)",
điểm đánh "a.", danh sách nơi nhận cùng cỡ chữ với lời văn. Test khẳng định
đúng những gì phần mềm HỨA sẽ sửa — và cũng khẳng định vài thứ nó hứa KHÔNG
đụng tới.
"""
import io

import pytest
from docx import Document
from docx.shared import Mm, Pt

from backend.services.vb_format import ap_dung, bien_doi, nhan_dien, quy_chuan
from backend.services.vb_format.chuan_hoa import chuan_hoa


# ── Dựng văn bản mẫu ─────────────────────────────────────────────────────────
def _van_ban_sai() -> bytes:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Mm(25)
    s.top_margin = s.bottom_margin = Mm(25)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for dong in [
        "NGÂN HÀNG NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN VIỆT NAM",
        "CHI NHÁNH HÀ NỘI",
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "Số: 05/QĐ-NHNo.HN-KTNQ",
        "Hà Nội, ngày 05 tháng 01 năm 2026",
        "QUYẾT ĐỊNH",
        "Về việc điều động cán bộ",
        "Căn cứ Quy chế số 616/QC-HĐTV-PC ngày 30/9/2022 của Hội đồng thành viên;",
        "Điều 1. Phạm vi điều chỉnh",
        "1) Quy định này áp dụng cho toàn hệ thống. quyết định có hiệu lực từ ngày ký.",
        "a. Các đơn vị tại trụ sở chính thực hiện theo khoản 2 điều 5 của quy chế.",
        "• Phòng Kế toán Ngân quỹ chịu trách nhiệm thi hành.",
        "Điều 2. Trách nhiệm thi hành",
        "2) nhà nước giao Tổng Giám đốc tổ chức thực hiện.",
        "Nơi nhận:",
        "- Như trên;",
        "- Ban kiểm soát;",
        "- Lưu: VT, PC.",
        "GIÁM ĐỐC",
        "Nguyễn Văn A",
    ]:
        doc.add_paragraph(dong)

    ra = io.BytesIO()
    doc.save(ra)
    return ra.getvalue()


@pytest.fixture(scope="module")
def ket_qua():
    du_lieu, bao_cao = chuan_hoa(_van_ban_sai())
    doc = Document(io.BytesIO(du_lieu))
    doan = ap_dung.duyet_doan(doc)
    return doc, [p for p, _ in doan], bao_cao


def _tim(dsach, mo_dau: str):
    for p in dsach:
        if p.text.strip().startswith(mo_dau):
            return p
    raise AssertionError(f"Không tìm thấy đoạn bắt đầu bằng {mo_dau!r}")


# ── Nhận diện thành phần thể thức ────────────────────────────────────────────
def test_nhan_dien_du_thanh_phan_the_thuc():
    doc = Document(io.BytesIO(_van_ban_sai()))
    khoi = ap_dung.duyet_doan(doc)
    ma = nhan_dien.phan_loai([(p.text, tb) for p, tb in khoi])
    theo_text = {p.text.strip(): m for (p, _), m in zip(khoi, ma)}

    assert theo_text["CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"] == "quoc_hieu"
    assert theo_text["Độc lập - Tự do - Hạnh phúc"] == "tieu_ngu"
    assert theo_text["CHI NHÁNH HÀ NỘI"] == "ten_dv_ban_hanh"
    assert theo_text["NGÂN HÀNG NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN VIỆT NAM"] == "ten_dv_chu_quan"
    assert theo_text["Số: 05/QĐ-NHNo.HN-KTNQ"] == "so_ky_hieu"
    assert theo_text["Hà Nội, ngày 05 tháng 01 năm 2026"] == "dia_danh_ngay"
    assert theo_text["QUYẾT ĐỊNH"] == "ten_loai"
    assert theo_text["Về việc điều động cán bộ"] == "trich_yeu"
    assert theo_text["Căn cứ Quy chế số 616/QC-HĐTV-PC ngày 30/9/2022 của Hội đồng thành viên;"] == "can_cu"
    assert theo_text["Điều 1. Phạm vi điều chỉnh"] == "dieu"
    assert theo_text["Nơi nhận:"] == "noi_nhan_tieu_de"
    assert theo_text["- Như trên;"] == "noi_nhan_ds"
    assert theo_text["- Lưu: VT, PC."] == "noi_nhan_ds"
    assert theo_text["GIÁM ĐỐC"] == "quyen_han_chuc_vu"
    assert theo_text["Nguyễn Văn A"] == "ho_ten_nguoi_ky"


def test_quoc_hieu_trong_cau_van_khong_bi_nham():
    """Câu NÓI VỀ Quốc hiệu không được nhận là Quốc hiệu rồi bị ép in hoa."""
    cau = '1. Quốc hiệu “CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM”: được trình bày bằng chữ in hoa.'
    ma = nhan_dien.phan_loai([(cau, False)])
    assert ma[0] != "quoc_hieu"


def test_la_in_hoa_nhan_dung_chu_hoa_co_dau():
    # "QUY ĐỊNH CHUNG" chứa "Ị" (U+1ECA) — từng bị dải [ạ-ỹ] hiểu là chữ thường
    assert nhan_dien.la_in_hoa("QUY ĐỊNH CHUNG")
    assert nhan_dien.la_in_hoa("TỔNG GIÁM ĐỐC")
    assert not nhan_dien.la_in_hoa("Quy định chung")


# ── Định dạng ────────────────────────────────────────────────────────────────
def test_le_trang_ve_dung_quy_dinh(ket_qua):
    doc, _, _ = ket_qua
    s = doc.sections[0]
    assert round(s.left_margin.mm) == 30
    assert round(s.right_margin.mm) == 20
    assert round(s.top_margin.mm) == 20
    assert round(s.bottom_margin.mm) == 20


def test_phong_chu_va_co_chu_theo_thanh_phan(ket_qua):
    _, doan, _ = ket_qua
    quoc_hieu = _tim(doan, "CỘNG HÒA")
    assert quoc_hieu.runs[0].font.name == "Times New Roman"
    # 12 chứ không phải 13: đếm trên cả 18 mẫu Phụ lục V thì Quốc hiệu là cỡ 12
    # ở 17 mẫu. Dải Phụ lục III ghi "12 - 13" — lấy cận trên làm dòng tên đơn vị
    # dài tràn cột và đẩy chữ "NAM" xuống một dòng riêng.
    assert quoc_hieu.runs[0].font.size.pt == 12
    assert _tim(doan, "CHI NHÁNH HÀ NỘI").runs[0].font.size.pt == 12

    # Danh sách nơi nhận cỡ 11, từ "Nơi nhận:" cỡ 12 nghiêng đậm (Điều 15.4.b).
    # Đọc cỡ chữ ĐANG CÓ HIỆU LỰC chứ không đọc `run.font.size`: văn bản mẫu đặt
    # style Normal 11pt, mà quy chuẩn cũng đòi 11 — đúng rồi thì phần mềm KHÔNG ghi đè,
    # nên trên run không có giá trị nào cả. Đó chính là hành vi mong muốn.
    o_noi_nhan = _tim(doan, "- Như trên")
    assert ap_dung._hieu_luc_run(o_noi_nhan.runs[0], o_noi_nhan, "size").pt == 11
    tieu_de = _tim(doan, "Nơi nhận:")
    assert tieu_de.runs[0].font.size.pt == 12
    assert tieu_de.runs[0].font.italic is True
    assert tieu_de.runs[0].font.bold is True

    # Lời văn cỡ 14
    assert _tim(doan, "Điều 1.").runs[0].font.size.pt == 14


def test_phong_chu_dat_ca_nhanh_complex_script(ket_qua):
    """Chữ tiếng Việt có dấu hay rơi vào nhánh w:cs — bỏ sót là hai kiểu chữ
    trên cùng một dòng khi mở ở máy khác."""
    from docx.oxml.ns import qn
    _, doan, _ = ket_qua
    r = _tim(doan, "Điều 1.").runs[0]
    rFonts = r._element.rPr.rFonts
    for thuoc in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        assert rFonts.get(qn(thuoc)) == "Times New Roman"


# ── Sửa chữ ──────────────────────────────────────────────────────────────────
def test_chuan_hoa_danh_so_va_gach_dau_dong(ket_qua):
    _, doan, _ = ket_qua
    chu = [p.text.strip() for p in doan]
    assert any(t.startswith("1. Quy định này") for t in chu), "khoản «1)» phải thành «1.»"
    assert any(t.startswith("a) Các đơn vị") for t in chu), "điểm «a.» phải thành «a)»"
    assert any(t.startswith("- Phòng Kế toán") for t in chu), "«•» phải thành «- »"


def test_viet_hoa_dau_cau_va_tu_dien(ket_qua):
    _, doan, _ = ket_qua
    chu = "\n".join(p.text for p in doan)
    assert "Quyết định có hiệu lực" in chu          # sau dấu chấm → viết hoa
    assert "2. Nhà nước giao" in chu                 # đầu dòng + từ điển "Nhà nước"
    assert "Ban Kiểm soát" in chu                    # từ điển


def test_vien_dan_dieu_khoan_diem(ket_qua):
    _, doan, _ = ket_qua
    chu = "\n".join(p.text for p in doan)
    # Phụ lục IV mục V.7: Điều viết hoa, khoản viết thường
    assert "khoản 2 Điều 5" in chu


def test_ten_don_vi_giu_nguyen_in_hoa(ket_qua):
    """Từ điển có "Ngân hàng Nông nghiệp…" nhưng dòng tên đơn vị đang IN HOA
    toàn bộ — hạ nó xuống dạng từ điển là làm sai đúng chỗ quy định bắt đúng."""
    _, doan, _ = ket_qua
    assert _tim(doan, "NGÂN HÀNG").text.strip() == (
        "NGÂN HÀNG NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN VIỆT\u00a0NAM"), \
        "chỉ dấu cách giữa VIỆT và NAM được đổi thành dấu cách không ngắt"


def test_cum_tu_lien_dong_dung_dau_cach_khong_ngat(ket_qua):
    """Cụm chức danh phải dính liền bằng dấu cách KHÔNG NGẮT (U+00A0).

    Viết hẳn \\u00a0 chứ không dán ký tự thật vào chuỗi: dấu cách
    không ngắt nhìn y hệt dấu cách thường trên màn hình, dán thật thì người
    đọc test không biết nó đang kiểm tra cái gì, mà sửa nhầm một ký tự là
    test vẫn xanh.
    """
    _, doan, _ = ket_qua
    chu = "\n".join(p.text for p in doan)
    assert "Tổng\u00a0Giám\u00a0đốc" in chu
    assert "Tổng Giám đốc" not in chu, "dấu cách thường phải được thay hết"


# ── Đánh dấu ─────────────────────────────────────────────────────────────────
def test_co_danh_dau_va_dung_ba_mau(ket_qua):
    _, doan, _ = ket_qua
    mau = {str(r.font.highlight_color) for p in doan for r in p.runs
           if r.font.highlight_color is not None}
    assert mau, "phải có vùng được đánh dấu"
    assert any("YELLOW" in m for m in mau)
    assert any("BRIGHT_GREEN" in m for m in mau)


def test_tat_danh_dau_thi_khong_boi_mau():
    du_lieu, _ = chuan_hoa(_van_ban_sai(), {"danh_dau": {"bat": False}})
    doc = Document(io.BytesIO(du_lieu))
    assert not [r for p, _ in ap_dung.duyet_doan(doc) for r in p.runs
                if r.font.highlight_color is not None]


# ── Báo cáo ──────────────────────────────────────────────────────────────────
def test_bao_cao_co_thong_ke_va_nhat_ky(ket_qua):
    _, _, bc = ket_qua
    assert bc["thong_ke"]["tong_doan"] == 21
    assert 0 < bc["thong_ke"]["doan_da_sua"] <= 21
    assert any("lề trái" in m for m in bc["sua_chung"])
    assert any("phông chữ" in m for m in bc["sua_chung"])
    assert all({"stt", "ma", "nhan", "trich", "viec"} <= set(d) for d in bc["doan"])


def test_gia_tri_chung_vao_sua_chung_khong_lap_o_tung_doan(ket_qua):
    """Giãn dòng / cách đoạn CỦA LỜI VĂN nằm ở «sửa chung», không lặp từng đoạn.

    Khối thể thức đầu và cuối trang thì ngược lại: chúng khai giãn dòng riêng
    (dòng đơn, 0pt) nên phải hiện ở đúng đoạn đó và được bôi màu — đó là khác
    biệt của riêng đoạn, không phải luật áp cho cả văn bản.
    """
    _, _, bc = ket_qua
    moi_viec = [v for d in bc["doan"] for v in d["viec"]]

    assert "giãn dòng → 1,2" in bc["sua_chung"]
    assert "cách đoạn → 6 pt" in bc["sua_chung"]
    assert "giãn dòng → 1,2" not in moi_viec
    assert "cách đoạn → 6 pt" not in moi_viec
    assert not [v for v in moi_viec if v.startswith("phông chữ")]


def test_khoi_the_thuc_dau_trang_dung_dong_don_khong_cach_doan(ket_qua):
    """Điều 7.3 / 8.2: Quốc hiệu, Tiêu ngữ, tên đơn vị cách nhau DÒNG ĐƠN.

    Ép 1,2 và 6pt cho cả khối này là lỗi đã gặp: Tiêu ngữ bị đẩy xa Quốc hiệu,
    khối đầu trang cao gấp đôi mẫu Phụ lục V.
    """
    def _cach_doan(p) -> float:
        # Đọc giá trị ĐANG CÓ HIỆU LỰC: đúng sẵn thì phần mềm không ghi đè, khi
        # đó `paragraph_format.space_after` là None chứ không phải 0.
        v = ap_dung._hieu_luc_doan(p, "space_after")
        return 0.0 if v is None else v.pt

    _, doan, _ = ket_qua
    for mo_dau in ("CỘNG HÒA", "Độc lập", "CHI NHÁNH HÀ NỘI", "Số:", "QUYẾT ĐỊNH"):
        p = _tim(doan, mo_dau)
        assert ap_dung._hieu_luc_doan(p, "line_spacing") == 1.0, mo_dau
        assert _cach_doan(p) == 0, mo_dau

    # Lời văn vẫn theo giá trị chung
    p = _tim(doan, "Điều 1.")
    assert ap_dung._hieu_luc_doan(p, "line_spacing") == 1.2
    assert _cach_doan(p) == 6


# ── Giữ nguyên định dạng bên trong đoạn ──────────────────────────────────────
def test_sua_chu_khong_lam_mat_dinh_dang_giua_cau():
    """Sửa một chỗ trong câu không được xoá phần in đậm ở chỗ khác của cùng câu."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Căn cứ ")
    r_dam = p.add_run("Quyết định số 600/QĐ-HĐTV")
    r_dam.bold = True
    p.add_run(" của hội đồng thành viên. quyết định này có hiệu lực.")
    ra = io.BytesIO()
    doc.save(ra)

    du_lieu, _ = chuan_hoa(ra.getvalue())
    lai = Document(io.BytesIO(du_lieu))
    doan = lai.paragraphs[0]
    assert "Quyết định này có hiệu lực" in doan.text        # đã sửa viết hoa đầu câu
    assert any(r.bold and "600/QĐ-HĐTV" in r.text for r in doan.runs), \
        "phần in đậm giữa câu phải còn nguyên"


def test_ap_sua_text_vat_qua_nhieu_run():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("nhà ")
    p.add_run("nước")
    p.add_run(" giao.")
    ap_dung.ap_sua_text(p, [(0, 8, "Nhà nước")])
    assert p.text == "Nhà nước giao."


# ── Cấu hình ─────────────────────────────────────────────────────────────────
def test_hop_nhat_giu_mac_dinh_cho_khoa_thieu():
    cfg = quy_chuan.hop_nhat({"thanh_phan": {"noi_dung": {"co": 13}}})
    assert cfg["thanh_phan"]["noi_dung"]["co"] == 13
    assert cfg["thanh_phan"]["noi_dung"]["can"] == "justify"      # khoá không khai
    assert cfg["chung"]["phong_chu"] == "Times New Roman"          # cả nhóm không khai


def test_hop_nhat_bo_qua_khoa_la():
    cfg = quy_chuan.hop_nhat({"khong_ton_tai": {"a": 1}})
    assert "khong_ton_tai" not in cfg


def test_mac_dinh_tra_ban_sao_doc_lap():
    a = quy_chuan.mac_dinh()
    a["chung"]["phong_chu"] = "Arial"
    assert quy_chuan.mac_dinh()["chung"]["phong_chu"] == "Times New Roman"


def test_cau_hinh_co_chu_duoc_ton_trong():
    du_lieu, _ = chuan_hoa(_van_ban_sai(), {"thanh_phan": {"noi_dung": {"co": 13}}})
    doc = Document(io.BytesIO(du_lieu))
    doan = [p for p, _ in ap_dung.duyet_doan(doc)]
    assert _tim(doan, "- Phòng Kế toán").runs[0].font.size.pt == 13


# ── Viết tắt không bị hiểu là hết câu ────────────────────────────────────────
@pytest.mark.parametrize("cau, mong_doi", [
    ("Trụ sở tại TP. hà nội.", "Trụ sở tại TP. hà nội."),
    ("Gồm sổ sách, chứng từ v.v. các tài liệu khác.", "Gồm sổ sách, chứng từ v.v. các tài liệu khác."),
    ("Đơn vị thực hiện. đơn vị báo cáo.", "Đơn vị thực hiện. Đơn vị báo cáo."),
])
def test_viet_hoa_dau_cau_bo_qua_viet_tat(cau, mong_doi):
    sua = bien_doi.viet_hoa_dau_cau(cau)
    ra = cau
    for dau, cuoi, moi in sorted(sua, reverse=True):
        ra = ra[:dau] + moi + ra[cuoi:]
    assert ra == mong_doi


def test_danh_so_khong_dung_vao_so_ky_hieu():
    """«Số: 05/QĐ-…» mà đem chuẩn hoá theo luật khoản sẽ thành «Số. 05/QĐ-…»."""
    assert bien_doi.chuan_danh_so("Số: 05/QĐ-NHNo-PC", "so_ky_hieu",
                                  quy_chuan.mac_dinh()["danh_so"]) == []


# ── Danh sách tự động của Word ───────────────────────────────────────────────
def _van_ban_co_danh_sach() -> bytes:
    doc = Document()
    doc.add_paragraph("Điều 1. Trách nhiệm thi hành")
    doc.add_paragraph("Phòng Kế toán thực hiện.", style="List Bullet")
    doc.add_paragraph("Bước một", style="List Number")
    ra = io.BytesIO()
    doc.save(ra)
    return ra.getvalue()


def test_nhan_ra_danh_sach_tu_dong_khai_bao_tren_STYLE():
    """python-docx và nút style của Word đặt w:numPr trên STYLE, không trên đoạn.
    Chỉ đọc trên đoạn là bỏ sót, mà bỏ sót thì im lặng — không lỗi nào báo."""
    doc = Document(io.BytesIO(_van_ban_co_danh_sach()))
    kieu = {p.text: ap_dung._kieu_danh_so(doc, p) for p, _ in ap_dung.duyet_doan(doc)}
    assert kieu["Phòng Kế toán thực hiện."] == "bullet"
    assert kieu["Bước một"] == "so"
    assert kieu["Điều 1. Trách nhiệm thi hành"] is None


def test_bullet_tu_dong_thanh_gach_dau_dong_con_danh_so_thi_canh_bao():
    du_lieu, bao_cao = chuan_hoa(_van_ban_co_danh_sach())
    doan = [p for p, _ in ap_dung.duyet_doan(Document(io.BytesIO(du_lieu)))]
    chu = [p.text.strip() for p in doan]

    assert "- Phòng Kế toán thực hiện." in chu
    # Đổi style về Normal, KHÔNG xoá numPr của style dùng chung
    assert _tim(doan, "- Phòng Kế toán").style.name == "Normal"
    # Danh sách ĐÁNH SỐ giữ nguyên + có cảnh báo nói rõ vì sao
    assert "Bước một" in chu
    assert any("ĐÁNH SỐ tự động" in w for w in bao_cao["luu_y"])


# ── Tên loại ngoài danh sách Điều 3, ngày tháng để trống, Tiêu ngữ ──────────
def _de_cuong() -> bytes:
    """Khối đầu một đề cương kiểm tra — dựng theo đúng file người dùng gửi."""
    doc = Document()
    for t in [
        "NGÂN HÀNG NÔNG NGHIỆP",
        "VÀ PHÁT TRIỂN NÔNG THÔN VIỆT NAM",
        "TRUNG TÂM THANH TOÁN",
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập %s Tự do %s Hạnh phúc" % ("–", "–"),
        "Số:……../TTTT-KSNB",
        "Hà Nội, ngày      tháng      năm 2026",
        "ĐỀ CƯƠNG",
        "KIỂM TRA HOẠT ĐỘNG TẠI TRUNG TÂM THANH TOÁN",
        "I. MỤC ĐÍCH, YÊU CẦU",
    ]:
        doc.add_paragraph(t)
    ra = io.BytesIO()
    doc.save(ra)
    return ra.getvalue()


def _ma_theo_text(du_lieu: bytes) -> dict:
    doc = Document(io.BytesIO(du_lieu))
    khoi = ap_dung.duyet_doan(doc)
    ma = nhan_dien.phan_loai([(p.text, tb) for p, tb in khoi])
    return {p.text.strip(): m for (p, _), m in zip(khoi, ma)}


def test_ten_loai_ngoai_danh_sach_dieu_3_van_duoc_nhan():
    """Điều 3.2.aa cho phép "các loại văn bản… khác phù hợp với thực tiễn" nên
    danh sách tên loại không bao giờ đủ. "ĐỀ CƯƠNG" từng bị xếp thành lời văn
    rồi bị căn đều hai bên thay vì canh giữa."""
    ma = _ma_theo_text(_de_cuong())
    assert ma["ĐỀ CƯƠNG"] == "ten_loai"
    assert ma["KIỂM TRA HOẠT ĐỘNG TẠI TRUNG TÂM THANH TOÁN"] == "trich_yeu"

    doan = [p for p, _ in ap_dung.duyet_doan(Document(io.BytesIO(chuan_hoa(_de_cuong())[0])))]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    assert _tim(doan, "ĐỀ CƯƠNG").paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_ngay_thang_de_trong_van_la_dia_danh_ngay():
    """Dự thảo trình ký và mọi mẫu Phụ lục V đều ghi "ngày  tháng  năm 2026"."""
    ma = _ma_theo_text(_de_cuong())
    assert ma["Hà Nội, ngày      tháng      năm 2026"] == "dia_danh_ngay"


def test_cau_vien_dan_ngay_thang_khong_bi_nham_la_dia_danh():
    cau = ("Căn cứ Quyết định số 05/QĐ-NHNo ngày 05 tháng 01 năm 2026 của "
           "Tổng Giám đốc về việc ban hành quy chế;")
    assert nhan_dien.phan_loai([(cau, False)])[0] == "can_cu"


def test_chuan_hoa_tieu_ngu_ve_dung_gach_noi_mot_dau_cach():
    """Điều 7.2: giữa các cụm từ có gạch NỐI (-), có cách chữ."""
    for goc in ("Độc lập – Tự do – Hạnh phúc",
                "Độc lập — Tự  do — Hạnh phúc",
                "Độc  lập  -  Tự do  -  Hạnh phúc"):
        sua = bien_doi.chuan_tieu_ngu(goc)
        ra = goc
        for dau, cuoi, moi in sorted(sua, reverse=True):
            ra = ra[:dau] + moi + ra[cuoi:]
        assert ra == "Độc lập - Tự do - Hạnh phúc", goc
    # Đúng rồi thì không sửa — không bôi màu một đoạn không đổi gì
    assert bien_doi.chuan_tieu_ngu("Độc lập - Tự do - Hạnh phúc") == []


def test_tieu_ngu_trong_van_ban_that_duoc_chuan_hoa():
    doan = [p for p, _ in ap_dung.duyet_doan(Document(io.BytesIO(chuan_hoa(_de_cuong())[0])))]
    assert _tim(doan, "Độc lập").text.strip() == "Độc lập - Tự do - Hạnh phúc"


def test_gian_dong_mac_dinh_la_1_2_theo_van_ban_979():
    """1,5 là cận TRÊN của dải Điều 12.6; lời văn của chính QĐ 979 dùng 1,2."""
    assert quy_chuan.mac_dinh()["chung"]["gian_dong"] == 1.2


# ── Công văn: khối đầu dựng bằng bảng, Kính gửi, ngắt dòng thẩm mỹ ───────────
def _cong_van() -> bytes:
    """Dựng theo đúng file người dùng gửi: khối đầu là bảng hai cột, mọi đoạn
    đặt Spacing Before 7pt / After 7pt."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(13)

    t = doc.add_table(rows=1, cols=2)
    trai, phai = t.rows[0].cells
    trai.text = ""
    phai.text = ""
    for o, dong in (
        (trai, ["NGÂN HÀNG NÔNG NGHIỆP",
                "VÀ PHÁT TRIỂN NÔNG THÔN VIỆT NAM",
                "Số:            /NHNo-TTTT",
                "V/v Thông báo thay đổi tên/địa chỉ đăng ký",
                "trên hệ thống SWIFT"]),
        (phai, ["CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM",
                "Độc lập - Tự do - Hạnh phúc",
                "Hà Nội, ngày 28 tháng 8 năm 2026"]),
    ):
        for txt in dong:
            p = o.add_paragraph(txt)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(7)

    for txt in ["Kính gửi: Giám đốc Agribank Quảng Ninh",
                "Ngày 20/8/2026, Trung tâm Thanh toán nhận được công văn.",
                "Trân trọng./."]:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(7)

    ra = io.BytesIO()
    doc.save(ra)
    return ra.getvalue()


@pytest.fixture(scope="module")
def cong_van():
    du_lieu, bao_cao = chuan_hoa(_cong_van())
    doan = [p for p, _ in ap_dung.duyet_doan(Document(io.BytesIO(du_lieu)))]
    return doan, bao_cao


def test_ten_nuoc_khong_bi_tach_lam_hai_dong(cong_van):
    """Gặp thật: dòng tên đơn vị bị ngắt thành "…NÔNG THÔN VIỆT" rồi "NAM" nằm
    một mình ở dòng thứ ba."""
    doan, _ = cong_van
    assert _tim(doan, "VÀ PHÁT TRIỂN").text.strip().endswith("VIỆT NAM")
    assert _tim(doan, "CỘNG HO").text.strip().endswith("VIỆT NAM")


def test_co_chu_khoi_dau_theo_dung_mau_979(cong_van):
    """Quốc hiệu và tên đơn vị cỡ 12 (đếm trên 18 mẫu Phụ lục V), trích yếu
    công văn cỡ 12, số ký hiệu 13, Tiêu ngữ 13."""
    doan, _ = cong_van

    def co(mo_dau: str) -> float:
        # Đọc cỡ ĐANG CÓ HIỆU LỰC: đúng sẵn thì phần mềm không ghi đè lên run
        p = _tim(doan, mo_dau)
        return ap_dung._hieu_luc_run(p.runs[0], p, "size").pt
    assert co("CỘNG HO") == 12
    assert co("NGÂN HÀNG") == 12
    assert co("VÀ PHÁT TRIỂN") == 12
    assert co("V/v") == 12
    assert co("Số:") == 13
    assert co("Độc lập") == 13


def test_kinh_gui_mot_noi_thi_canh_giua(cong_van):
    """Điều 15.4.a: gửi MỘT nơi thì "Kính gửi" và tên đơn vị trên cùng một
    dòng — mẫu 06 và 09 của Phụ lục V canh giữa dòng đó."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doan, _ = cong_van
    assert _tim(doan, "Kính gửi").paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_kinh_gui_nhieu_noi_thi_sat_trai():
    """Chỉ có chữ "Kính gửi:" rồi liệt kê xuống dòng → canh giữa là sai."""
    ma = nhan_dien.phan_loai([("Kính gửi:", False)])
    assert ma[0] == "kinh_gui_ds"
    assert nhan_dien.phan_loai([("Kính gửi: Ngân hàng Nhà nước Việt Nam", False)])[0] \
        == "kinh_gui"


def test_khoi_dau_ve_spacing_0_0_ke_ca_o_bang_khong_nhan_ra(cong_van):
    """Ô bảng không khớp thành phần nào VẪN phải về 0/0 — bỏ sót thì khối đầu
    vẫn giãn dù mọi đoạn nhận ra đều đã về 0."""
    doan, _ = cong_van

    def _sp(p):
        tr = ap_dung._hieu_luc_doan(p, "space_before")
        sa = ap_dung._hieu_luc_doan(p, "space_after")
        return (0.0 if tr is None else tr.pt, 0.0 if sa is None else sa.pt)

    for mo_dau in ("NGÂN HÀNG", "CỘNG HO", "Độc lập", "Số:", "V/v",
                   "trên hệ thống", "Hà Nội,", "Kính gửi"):
        assert _sp(_tim(doan, mo_dau)) == (0.0, 0.0), mo_dau


def test_loi_van_bo_khoang_truoc_giu_khoang_sau(cong_van):
    """Khoảng cách giữa hai đoạn = after của đoạn trên + before của đoạn dưới.
    7+7 cho ra 14pt mà hộp Paragraph chỉ hiện hai số 7 — đưa before về 0 để
    chỉ còn một nguồn. `after` giữ nguyên 7 vì Điều 12.6 chỉ nêu mức tối thiểu
    6pt, hạ xuống là sửa thứ không sai."""
    doan, _ = cong_van
    p = _tim(doan, "Ngày 20/8/2026")
    assert ap_dung._hieu_luc_doan(p, "space_before").pt == 0
    assert ap_dung._hieu_luc_doan(p, "space_after").pt == 7


def test_xuong_dong_tham_my_khong_bi_viet_hoa(cong_van):
    """Ô trích yếu xuống dòng cho cân ô, không phải hết câu — "trên hệ thống
    SWIFT" phải giữ chữ thường."""
    doan, _ = cong_van
    assert _tim(doan, "trên hệ thống").text.strip().startswith("trên")


def test_van_viet_hoa_khi_doan_truoc_da_ket_cau():
    """Ngược lại: đoạn trước kết thúc bằng dấu chấm thì đây là câu mới."""
    assert bien_doi.cho_phep_hoa_dau_doan("noi_dung", "Đơn vị thực hiện.")
    assert not bien_doi.cho_phep_hoa_dau_doan("noi_dung", "V/v Thông báo thay đổi")
    # Thành phần không phải lời văn thì không bao giờ áp luật viết hoa đầu dòng
    assert not bien_doi.cho_phep_hoa_dau_doan("trich_yeu_cong_van", "Đơn vị thực hiện.")
    assert not bien_doi.cho_phep_hoa_dau_doan("noi_nhan_ds", None)
